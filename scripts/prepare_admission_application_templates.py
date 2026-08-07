from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


PLACEHOLDERS = {
    "full": "{{ side_1_full_name }}",
    "full_genitive": "{{ side_1_full_name_genitive }}",
    "iin": "{{ side_1_iin_bin }}",
    "basis_ru": "{{ admission_basis_ru }}",
    "basis_kk": "{{ admission_basis_kk }}",
    "pg_ru": "{{ program_group_code }} {{ program_group_name_ru }}",
    "pg_kk": "{{ program_group_code }} {{ program_group_name_kk }}",
    "p_ru": "{{ program_code }} {{ program_name_ru }}",
    "p_kk": "{{ program_code }} {{ program_name_kk }}",
    "fac_ru": "{{ program_faculty_ru }}",
    "fac_kk": "{{ program_faculty_kk }}",
    "form_ru": "{{ study_form_ru }}",
    "form_kk": "{{ study_form_kk }}",
    "lang_ru": "{{ study_language_ru }}",
    "lang_kk": "{{ study_language_kk }}",
    "birth_ru": "{{ birth_date_text_ru }}",
    "birth_kk": "{{ birth_date_text_kk }}",
    "idn": "{{ identity_document_number }}",
    "id_issue_ru": "{{ identity_document_issue_date_ru }}",
    "id_issue_kk": "{{ identity_document_issue_date_kk }}",
    "issuer_ru": "{{ identity_document_issuer_ru }}",
    "issuer_kk": "{{ identity_document_issuer_kk }}",
    "gender_ru": "{{ gender_ru }}",
    "gender_kk": "{{ gender_kk }}",
    "cit_ru": "{{ citizenship_ru }}",
    "cit_kk": "{{ citizenship_kk }}",
    "nat_ru": "{{ nationality_ru }}",
    "nat_kk": "{{ nationality_kk }}",
    "grad": "{{ graduation_year }}",
    "prev_ru": "{{ previous_education_ru }}",
    "prev_kk": "{{ previous_education_kk }}",
    "edu_type_ru": "{{ education_document_type_ru }}",
    "edu_type_kk": "{{ education_document_type_kk }}",
    "edu_series": "{{ education_document_series }}",
    "edu_number": "{{ education_document_number }}",
    "edu_issue": "{{ education_document_issue_date }}",
    "dist_ru": "{{ distinction_award_ru }}",
    "dist_kk": "{{ distinction_award_kk }}",
    "oly_sub_ru": "{{ olympiad_subject_ru }}",
    "oly_sub_kk": "{{ olympiad_subject_kk }}",
    "oly_deg_ru": "{{ olympiad_degree_ru }}",
    "oly_deg_kk": "{{ olympiad_degree_kk }}",
    "cert": "{{ certificate_score }}",
    "avg": "{{ average_grade }}",
    "quota_ru": "{{ admission_quota_ru }}",
    "quota_kk": "{{ admission_quota_kk }}",
    "father": "{{ father_full_name }}",
    "father_phone": "{{ father_phone }}",
    "father_work": "{{ father_work_place }}",
    "father_position": "{{ father_position }}",
    "mother": "{{ mother_full_name }}",
    "mother_phone": "{{ mother_phone }}",
    "mother_work": "{{ mother_work_place }}",
    "mother_position": "{{ mother_position }}",
    "email": "{{ side_1_email }}",
    "phone": "{{ side_1_phone }}",
    "almaty": "{{ almaty_address }}",
    "addr": "{{ student_address }}",
    "foreign_ru": "{{ foreign_language_ru }}",
    "foreign_kk": "{{ foreign_language_kk }}",
    "dorm_ru": "{{ dormitory_need_ru }}",
    "dorm_kk": "{{ dormitory_need_kk }}",
    "date_ru": "{{ contract_date_text_ru }}",
    "date_kk": "{{ contract_date_text_kk }}",
    "applicant_signature": "{{ applicant_signature_full_name }}",
    "secretary": "{{ technical_secretary_full_name }}",
    "dean": "{{ dean_full_name }}",
}


def set_paragraph_text(paragraph, text):
    text = tidy_placeholder_spacing(text)
    text = normalize_kazakh_letters(text)
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return

    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def set_paragraph_appended_text(paragraph, text):
    text = tidy_placeholder_spacing(text)
    text = normalize_kazakh_letters(text)
    run = paragraph.add_run(text)

    if paragraph.runs:
        source_run = paragraph.runs[0]
        run.bold = source_run.bold
        run.italic = source_run.italic
        run.underline = source_run.underline
        run.font.size = source_run.font.size

    return run


def put_text_on_line_and_set_caption(
    doc,
    *,
    line_paragraph_index,
    caption_paragraph_index,
    value,
    caption="",
):
    line_paragraph = doc.paragraphs[line_paragraph_index]
    line_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_appended_text(line_paragraph, value)
    set_paragraph_text(doc.paragraphs[caption_paragraph_index], caption)


def put_text_on_line_before_caption(doc, *, caption_contains, value):
    for caption_index, paragraph in enumerate(doc.paragraphs):
        if caption_contains not in paragraph.text:
            continue

        for line_index in range(caption_index - 1, -1, -1):
            candidate = doc.paragraphs[line_index]

            if candidate.text.strip():
                break

            if candidate.runs:
                candidate.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_appended_text(candidate, value)
                return

        return


def replace_tabs(text, replacements):
    parts = text.split("\t")
    result = [parts[0]]
    for index, part in enumerate(parts[1:]):
        result.append(replacements[index] if index < len(replacements) else "")
        result.append(part)
    return "".join(result)


def append_value(text, placeholder):
    if "\t" in text:
        return replace_tabs(text, [placeholder])
    return text.rstrip() + " " + placeholder


def prepend_value(text, placeholder):
    return placeholder + " " + text.lstrip()


def replace_date_signature_line(text, placeholder):
    parts = text.split("\t")
    if len(parts) >= 5:
        prefix = parts[0].split("\u00ab", 1)[0].strip()
        tail = " ".join(part.strip() for part in parts[4:] if part.strip())
        return " ".join(part for part in [prefix, placeholder, tail] if part)

    return placeholder


def tidy_placeholder_spacing(text):
    text = re.sub(r"(?<![\s\xab]){{", r" {{", text)
    text = re.sub(r"}}(?![\s,.;:\xbb)])", r"}} ", text)
    text = re.sub(r"(?<=}})_+", "", text)
    text = re.sub(r"_+(?={{)", "", text)
    text = re.sub(r"\s+_+(?=\s|$)", "", text)
    text = re.sub(r"[ ]{2,}", " ", text)
    text = re.sub(r"[ ]+([,.;:])", r"\1", text)
    text = text.replace(" \t", "\t").replace("\t ", "\t")
    return text.strip()


def normalize_kazakh_letters(text):
    return text.replace("Ə", "Ә").replace("ə", "ә")


def education_document_ru():
    return (
        f"В {PLACEHOLDERS['grad']} году окончил(а) {PLACEHOLDERS['prev_ru']}, "
        f"документ об образовании: {PLACEHOLDERS['edu_type_ru']}, "
        f"серия {PLACEHOLDERS['edu_series']}, № {PLACEHOLDERS['edu_number']}, "
        f"дата выдачи {PLACEHOLDERS['edu_issue']}"
    )


def education_document_kk():
    return (
        f"{PLACEHOLDERS['grad']} жылы {PLACEHOLDERS['prev_kk']} оқу орнын, "
        f"білімі туралы құжат: {PLACEHOLDERS['edu_type_kk']}, "
        f"сериясы {PLACEHOLDERS['edu_series']}, № {PLACEHOLDERS['edu_number']}, "
        f"берілген күні {PLACEHOLDERS['edu_issue']}"
    )


def parent_ru(parent_label, full_name_key, phone_key, work_key, position_key):
    return (
        f"{parent_label} {PLACEHOLDERS[full_name_key]}, "
        f"тел. {PLACEHOLDERS[phone_key]}, "
        f"место работы {PLACEHOLDERS[work_key]}, "
        f"должность {PLACEHOLDERS[position_key]}"
    )


def parent_kk(parent_label, full_name_key, phone_key, work_key, position_key):
    return (
        f"{parent_label} {PLACEHOLDERS[full_name_key]}, "
        f"тел. {PLACEHOLDERS[phone_key]}, "
        f"жұмыс орны {PLACEHOLDERS[work_key]}, "
        f"лауазымы {PLACEHOLDERS[position_key]}"
    )


def signature_ru():
    return (
        f"Дата {PLACEHOLDERS['date_ru']} Личная подпись "
        f"{PLACEHOLDERS['applicant_signature']} Подтверждаю достоверность внесенных данных"
    )


def secretary_dean_ru():
    return (
        f"Технический секретарь {PLACEHOLDERS['secretary']} / "
        f"Декан {PLACEHOLDERS['dean']}"
    )


def signature_kk():
    return (
        f"Күні {PLACEHOLDERS['date_kk']} Өтініш иесінің қолы "
        f"{PLACEHOLDERS['applicant_signature']}"
    )


def secretary_dean_kk():
    return (
        f"Техникалық хатшы {PLACEHOLDERS['secretary']} / "
        f"Декан {PLACEHOLDERS['dean']}"
    )


def nonempty_paragraphs(doc):
    return [paragraph for paragraph in doc.paragraphs if paragraph.text.strip()]


def iter_all_paragraphs(doc):
    yield from doc.paragraphs

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def normalize_document_text(doc):
    for paragraph in iter_all_paragraphs(doc):
        relax_paragraph_spacing(paragraph)
        for run in paragraph.runs:
            normalized = normalize_kazakh_letters(run.text)
            if normalized != run.text:
                run.text = normalized


def relax_paragraph_spacing(paragraph):
    paragraph_format = paragraph.paragraph_format
    if paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY:
        paragraph_format.line_spacing = 1.0
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def is_filler_paragraph(paragraph):
    text = paragraph.text.strip()
    return not text or re.fullmatch(r"_+", text) is not None


def remove_filler_between_markers(doc, start_marker, end_marker):
    paragraphs = doc.paragraphs
    start_index = next(
        (index for index, paragraph in enumerate(paragraphs) if start_marker in paragraph.text),
        None,
    )
    end_index = next(
        (index for index, paragraph in enumerate(paragraphs) if end_marker in paragraph.text),
        None,
    )
    if start_index is None or end_index is None or end_index <= start_index + 1:
        return

    for paragraph in reversed(paragraphs[start_index + 1:end_index]):
        if is_filler_paragraph(paragraph):
            delete_paragraph(paragraph)


def compact_bachelor_parent_blocks(doc):
    marker_pairs = [
        ("Отец {{ father_full_name }}", "Мать {{ mother_full_name }}"),
        ("Мать {{ mother_full_name }}", "E-mail абитуриента"),
        ("Әкем {{ father_full_name }}", "Анам {{ mother_full_name }}"),
        ("Анам {{ mother_full_name }}", "Талапкердің e-mail"),
    ]
    for start_marker, end_marker in marker_pairs:
        remove_filler_between_markers(doc, start_marker, end_marker)


def apply_edits(doc, edits):
    paragraphs = nonempty_paragraphs(doc)
    for paragraph_number, edit in edits.items():
        paragraph = paragraphs[paragraph_number - 1]
        set_paragraph_text(paragraph, edit(paragraph.text))


def prepare_bachelor(doc):
    apply_edits(doc, {
        3: lambda text: replace_tabs(text, [PLACEHOLDERS["full_genitive"]]),
        6: lambda text: append_value(text, PLACEHOLDERS["basis_ru"]),
        8: lambda text: replace_tabs(text, [PLACEHOLDERS["pg_ru"]]),
        10: lambda text: replace_tabs(text, [PLACEHOLDERS["p_ru"]]),
        12: lambda text: replace_tabs(text, [PLACEHOLDERS["fac_ru"]]),
        13: lambda text: replace_tabs(text, [PLACEHOLDERS["form_ru"], PLACEHOLDERS["lang_ru"]]),
        16: lambda text: (
            f"1. Дата рождения «{PLACEHOLDERS['birth_ru']}» года. "
            f"2. № удостоверения {PLACEHOLDERS['idn']}, ИИН {PLACEHOLDERS['iin']}, "
            f"выдано {PLACEHOLDERS['issuer_ru']}, дата выдачи {PLACEHOLDERS['id_issue_ru']}"
        ),
        17: lambda text: replace_tabs(text, [PLACEHOLDERS["gender_ru"], PLACEHOLDERS["cit_ru"], PLACEHOLDERS["nat_ru"]]),
        18: lambda text: education_document_ru(),
        20: lambda text: f"Награда / отличие: {PLACEHOLDERS['dist_ru']}",
        22: lambda text: f"Олимпиада / соревнование: {PLACEHOLDERS['oly_sub_ru']} {PLACEHOLDERS['oly_deg_ru']}",
        23: lambda text: "",
        24: lambda text: replace_tabs(text, [PLACEHOLDERS["cert"], PLACEHOLDERS["avg"]]),
        25: lambda text: f"Квота: {PLACEHOLDERS['quota_ru']}",
        27: lambda text: parent_ru("Отец", "father", "father_phone", "father_work", "father_position"),
        29: lambda text: parent_ru("Мать", "mother", "mother_phone", "mother_work", "mother_position"),
        30: lambda text: replace_tabs(text, [PLACEHOLDERS["email"]]),
        31: lambda text: replace_tabs(text, [PLACEHOLDERS["almaty"] + " " + PLACEHOLDERS["phone"]]),
        32: lambda text: replace_tabs(text, [PLACEHOLDERS["addr"]]),
        33: lambda text: replace_tabs(text, [PLACEHOLDERS["foreign_ru"]]),
        34: lambda text: f"Общежитие: {PLACEHOLDERS['dorm_ru']}",
        36: lambda text: signature_ru(),
        37: lambda text: secretary_dean_ru(),
        39: lambda text: "(өтініш берушінің аты-жөні)",
        41: lambda text: replace_tabs(text, [PLACEHOLDERS["basis_kk"]]),
        43: lambda text: replace_tabs(text, [PLACEHOLDERS["pg_kk"]]),
        45: lambda text: replace_tabs(text, [PLACEHOLDERS["p_kk"]]),
        47: lambda text: replace_tabs(text, [PLACEHOLDERS["fac_kk"]]),
        49: lambda text: replace_tabs(text, [PLACEHOLDERS["form_kk"], PLACEHOLDERS["lang_kk"]]),
        52: lambda text: (
            f"1. Туған күні «{PLACEHOLDERS['birth_kk']}» ж. "
            f"2. Жеке куәлігім № {PLACEHOLDERS['idn']}, ЖСН {PLACEHOLDERS['iin']}, "
            f"кіммен берілген {PLACEHOLDERS['issuer_kk']}, берілген күні {PLACEHOLDERS['id_issue_kk']}"
        ),
        53: lambda text: replace_tabs(text, [PLACEHOLDERS["gender_kk"], PLACEHOLDERS["cit_kk"], PLACEHOLDERS["nat_kk"]]),
        54: lambda text: education_document_kk(),
        57: lambda text: f"Марапат / ерекше белгі: {PLACEHOLDERS['dist_kk']}",
        59: lambda text: f"Олимпиада / жарыс: {PLACEHOLDERS['oly_sub_kk']} {PLACEHOLDERS['oly_deg_kk']}",
        60: lambda text: replace_tabs(text, [PLACEHOLDERS["cert"], PLACEHOLDERS["avg"]]),
        61: lambda text: f"Квота: {PLACEHOLDERS['quota_kk']}",
        62: lambda text: "",
        64: lambda text: parent_kk("Әкем", "father", "father_phone", "father_work", "father_position"),
        65: lambda text: parent_kk("Анам", "mother", "mother_phone", "mother_work", "mother_position"),
        66: lambda text: replace_tabs(text, [PLACEHOLDERS["email"]]),
        67: lambda text: replace_tabs(text, [PLACEHOLDERS["almaty"] + " " + PLACEHOLDERS["phone"]]),
        68: lambda text: replace_tabs(text, [PLACEHOLDERS["addr"]]),
        69: lambda text: replace_tabs(text, [PLACEHOLDERS["foreign_kk"]]),
        70: lambda text: f"Жатақхана: {PLACEHOLDERS['dorm_kk']}",
        72: lambda text: signature_kk(),
        74: lambda text: secretary_dean_kk(),
    })
    put_text_on_line_before_caption(
        doc,
        caption_contains="өтініш беруш",
        value=PLACEHOLDERS["full"],
    )
    compact_bachelor_parent_blocks(doc)


def prepare_master_doctoral(doc):
    apply_edits(doc, {
        3: lambda text: replace_tabs(text, [PLACEHOLDERS["full_genitive"]]),
        6: lambda text: replace_tabs(text, [PLACEHOLDERS["basis_ru"]]),
        8: lambda text: replace_tabs(text, [PLACEHOLDERS["pg_ru"]]),
        10: lambda text: replace_tabs(text, [PLACEHOLDERS["p_ru"]]),
        12: lambda text: replace_tabs(text, [PLACEHOLDERS["fac_ru"], ""]),
        13: lambda text: replace_tabs(text, [PLACEHOLDERS["form_ru"], PLACEHOLDERS["lang_ru"]]),
        16: lambda text: replace_tabs(text, [PLACEHOLDERS["birth_ru"], " "]),
        17: lambda text: (
            f"№ удостоверения личности или паспорта {PLACEHOLDERS['idn']} "
            f"ИИН {PLACEHOLDERS['iin']}. Когда и кем выдан {PLACEHOLDERS['issuer_ru']}, "
            f"дата выдачи {PLACEHOLDERS['id_issue_ru']}"
        ),
        18: lambda text: replace_tabs(text, [PLACEHOLDERS["gender_ru"], PLACEHOLDERS["cit_ru"], PLACEHOLDERS["nat_ru"]]),
        19: lambda text: education_document_ru(),
        21: lambda text: replace_tabs(text, [PLACEHOLDERS["cert"], PLACEHOLDERS["avg"]]),
        23: lambda text: (
            parent_ru("Отец", "father", "father_phone", "father_work", "father_position")
            + "; "
            + parent_ru("Мать", "mother", "mother_phone", "mother_work", "mother_position")
        ),
        24: lambda text: replace_tabs(text, [PLACEHOLDERS["almaty"]]),
        25: lambda text: replace_tabs(text, [PLACEHOLDERS["addr"]]),
        26: lambda text: replace_tabs(text, [PLACEHOLDERS["email"] + " " + PLACEHOLDERS["phone"]]),
        27: lambda text: f"Иностранный язык {PLACEHOLDERS['foreign_ru']}. Общежитие: {PLACEHOLDERS['dorm_ru']}",
        29: lambda text: signature_ru(),
        30: lambda text: secretary_dean_ru(),
        32: lambda text: "",
        34: lambda text: replace_tabs(text, [PLACEHOLDERS["basis_kk"]]),
        36: lambda text: replace_tabs(text, [PLACEHOLDERS["pg_kk"]]),
        38: lambda text: replace_tabs(text, [PLACEHOLDERS["p_kk"]]),
        40: lambda text: replace_tabs(text, [PLACEHOLDERS["fac_kk"]]),
        41: lambda text: replace_tabs(text, [PLACEHOLDERS["form_kk"], PLACEHOLDERS["lang_kk"]]),
        44: lambda text: replace_tabs(text, [PLACEHOLDERS["birth_kk"], " "]),
        45: lambda text: (
            f"Жеке куәлігім № {PLACEHOLDERS['idn']} "
            f"ЖСН {PLACEHOLDERS['iin']}. Қашан және кіммен берілген {PLACEHOLDERS['issuer_kk']}, "
            f"берілген күні {PLACEHOLDERS['id_issue_kk']}"
        ),
        46: lambda text: replace_tabs(text, [PLACEHOLDERS["gender_kk"], PLACEHOLDERS["cit_kk"], PLACEHOLDERS["nat_kk"]]),
        47: lambda text: education_document_kk(),
        50: lambda text: replace_tabs(text, [PLACEHOLDERS["cert"], PLACEHOLDERS["avg"]]),
        52: lambda text: parent_kk("Әкем", "father", "father_phone", "father_work", "father_position"),
        53: lambda text: parent_kk("Анам", "mother", "mother_phone", "mother_work", "mother_position"),
        54: lambda text: replace_tabs(text, [PLACEHOLDERS["almaty"]]),
        55: lambda text: replace_tabs(text, [PLACEHOLDERS["addr"]]),
        56: lambda text: replace_tabs(text, [PLACEHOLDERS["email"] + " " + PLACEHOLDERS["phone"]]),
        57: lambda text: f"Шет тілі {PLACEHOLDERS['foreign_kk']}. Жатақхана: {PLACEHOLDERS['dorm_kk']}",
        59: lambda text: signature_kk(),
        61: lambda text: secretary_dean_kk(),
    })
    put_text_on_line_and_set_caption(
        doc,
        line_paragraph_index=43,
        caption_paragraph_index=44,
        value=PLACEHOLDERS["full"],
    )


def find_docx(directory, marker):
    for path in directory.glob("*.docx"):
        if marker in path.name:
            return path
    raise FileNotFoundError(f"DOCX file with marker {marker!r} was not found in {directory}")


def prepare(marker, prepare_func):
    downloads_dir = Path.home() / "Downloads"
    prepared_dir = Path("prepared_templates") / "admissions"
    source_path = find_docx(downloads_dir, marker)
    output_path = find_docx(prepared_dir, marker)

    doc = Document(source_path)
    prepare_func(doc)
    normalize_document_text(doc)
    doc.save(output_path)
    return output_path


def validate(path):
    doc = Document(path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    text += "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    variables = sorted(set(re.findall(r"{{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*}}", text)))
    return text.count("?"), variables


def main():
    outputs = [
        prepare("10000", prepare_bachelor),
        prepare("3000", prepare_master_doctoral),
    ]

    for output in outputs:
        question_marks, variables = validate(output)
        print(output)
        print(f"question_marks={question_marks}")
        print(f"variables={len(variables)}")
        print(", ".join(variables))


if __name__ == "__main__":
    main()
