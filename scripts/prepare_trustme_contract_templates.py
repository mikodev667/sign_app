from pathlib import Path
import re

from docx import Document


SOURCE_DIR = Path("tmp/trustme_docs")
OUTPUT_DIR = Path("prepared_templates/trustme_docs")


P = {
    "contract_number": "{{ contract_number }}",
    "date_kk": "{{ contract_date_text_kk }}",
    "date_ru": "{{ contract_date_text_ru }}",
    "date_en": "{{ contract_date_text_en }}",
    "student": "{{ side_1_full_name }}",
    "student_iin": "{{ side_1_iin_bin }}",
    "student_address": "{{ side_1_address }}",
    "student_phone": "{{ side_1_phone }}",
    "student_email": "{{ side_1_email }}",
    "student_citizenship": "{{ side_1_citizenship }}",
    "student_identity": "{{ side_1_identity_document }}",
    "customer": "{{ side_2_full_name }}",
    "customer_address": "{{ side_2_address }}",
    "customer_email": "{{ side_2_email }}",
    "customer_representative": "{{ side_2_representative_full_name }}",
    "customer_authority": "{{ side_2_authority_basis }}",
    "program_kk": "{{ program_code }} {{ program_name_kk }}",
    "program_ru": "{{ program_code }} {{ program_name_ru }}",
    "program_en": "{{ program_code }} {{ program_name_en }}",
    "faculty_kk": "{{ program_faculty_kk }}",
    "faculty_ru": "{{ program_faculty_ru }}",
    "duration_ru": "{{ program_duration_ru }}",
    "duration_en": "{{ program_duration_en }}",
    "qualification_ru": "{{ qualification_ru }}",
    "qualification_en": "{{ qualification_en }}",
    "tuition_kk": "{{ tuition_amount_full_kk }}",
    "tuition_ru": "{{ tuition_amount_full_ru }}",
    "tuition_en": "{{ tuition_amount_full_en }}",
    "year_1_kk": "{{ year_1_amount_full_kk }}",
    "year_2_kk": "{{ year_2_amount_full_kk }}",
    "year_3_kk": "{{ year_3_amount_full_kk }}",
    "year_1_ru": "{{ year_1_amount_full_ru }}",
    "year_2_ru": "{{ year_2_amount_full_ru }}",
    "year_3_ru": "{{ year_3_amount_full_ru }}",
    "vice": "{{ university_representative_full_name }}",
    "authority_number": "{{ university_authority_number }}",
    "authority_date_kk": "{{ university_authority_date_kk }}",
    "authority_date_ru": "{{ university_authority_date_ru }}",
    "authority_date_en": "{{ university_authority_date_en }}",
    "authority_year": "{{ university_authority_year }}",
}


def iter_paragraphs(parent):
    yield from parent.paragraphs
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def set_paragraph_text(paragraph, text):
    text = tidy(text)
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return

    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def tidy(text):
    text = re.sub(r"(?<![\s\u00ab]){{", r" {{", text)
    text = re.sub(r"}}(?![\s,.;:\u00bb)])", r"}} ", text)
    text = re.sub(r"(?<=}})_+", "", text)
    text = re.sub(r"_+(?={{)", "", text)
    text = re.sub(r"\s+_+(?=\s|$)", "", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"[ ]+([,.;:])", r"\1", text)
    return text.strip()


def replace_blanks(text, replacements):
    parts = re.split(r"(_{2,})", text)
    result = []
    index = 0

    for part in parts:
        if re.fullmatch(r"_{2,}", part):
            if index < len(replacements):
                result.append(replacements[index])
                index += 1
            continue
        result.append(part)

    return "".join(result)


def number_line(text):
    return replace_blanks(text, [P["contract_number"]])


def date_line(text, placeholder):
    if "Алматы" in text:
        prefix = text.split("«", 1)[0].strip()
        return f"{prefix} {placeholder}"
    if "Almaty" in text:
        return f"Almaty {placeholder}"
    return placeholder


def clear(_text):
    return ""


def student_line(text):
    return replace_blanks(text, [P["student"]])


def customer_line(text):
    return replace_blanks(text, [P["customer"]])


def customer_representative_line(text):
    return replace_blanks(text, [P["customer_representative"]])


def customer_authority_student_line(text):
    return replace_blanks(text, [P["customer_authority"], P["student"]])


def faculty_line(text, placeholder):
    return replace_blanks(text, [placeholder])


def program_line(text, placeholder):
    return replace_blanks(text, [placeholder])


def amount_line(text, placeholder):
    return replace_blanks(text, [placeholder])


def installment_line(year_number, tranche_number):
    amount = f"{{{{ year_{year_number}_tranche_{tranche_number}_amount }}}}"
    due_date = f"{{{{ year_{year_number}_tranche_{tranche_number}_due_date }}}}"

    def edit(text):
        if "—" in text:
            label = text.split("—", 1)[0].strip()
        elif "-" in text:
            label = text.split("-", 1)[0].strip()
        else:
            label = text.strip()

        if "тг" in text:
            currency = "тг."
        elif "теңге" in text:
            currency = "теңге"
        else:
            currency = "тенге"

        if "дейін" in text:
            return f"{label} — {amount} {currency} мөлшерінде, {due_date} дейін."

        return f"{label} — {amount} {currency} до {due_date}."

    return edit


def clean_orphan_underscores(text):
    placeholders = []

    def hold(match):
        placeholders.append(match.group(0))
        return f"@@PH{len(placeholders) - 1}@@"

    protected = re.sub(r"{{\s*[^}]+?\s*}}", hold, text)
    protected = protected.replace("—_", "—").replace("— _", "—")
    protected = re.sub(r"\s+_+\s*", " ", protected)
    protected = re.sub(r"_+\s*", "", protected)

    for index, placeholder in enumerate(placeholders):
        protected = protected.replace(f"@@PH{index}@@", placeholder)

    return protected


def address_line(text):
    return replace_blanks(text, [P["student_address"]])


def phone_line(text):
    return replace_blanks(text, [P["student_phone"]])


def email_line(text):
    return replace_blanks(text, [P["student_email"]])


def customer_address_line(text):
    return replace_blanks(text, [P["customer_address"]])


def customer_email_line(text):
    return replace_blanks(text, [P["customer_email"]])


def consent_line(text):
    return replace_blanks(text, [P["student"]])


def receipt_kk(_text):
    return f"Шарттың бір данасын алдым {P['student']} / {P['date_kk']} /"


def receipt_ru(_text):
    return f"Экземпляр Договора получил(-а) {P['student']} / {P['date_ru']} /"


def authority_year(text):
    return text.replace("202__", P["authority_year"]).replace("202_", P["authority_year"])


def intro_two_party_kk(text):
    text = authority_year(text)
    return replace_blanks(
        text,
        [
            P["authority_date_kk"],
            P["authority_number"],
            P["vice"],
            P["student"],
        ],
    )


def intro_two_party_ru(text):
    text = authority_year(text)
    return replace_blanks(
        text,
        [
            P["vice"],
            P["authority_number"],
            P["authority_date_ru"],
            P["student"],
        ],
    )


def intro_tripartite_kk(text):
    text = authority_year(text)
    return replace_blanks(
        text,
        [
            P["authority_date_kk"],
            P["authority_number"],
            P["vice"],
            P["student"],
        ],
    )


def intro_tripartite_ru(text):
    text = authority_year(text)
    return replace_blanks(
        text,
        [
            P["vice"],
            P["authority_number"],
            P["authority_date_ru"],
            P["student"],
        ],
    )


def intro_fibs_ru(text):
    text = text.replace("Галеевой А.К.", P["vice"])
    text = text.replace("№28", f"№ {P['authority_number']}")
    text = text.replace("№ 28", f"№ {P['authority_number']}")
    text = text.replace("11.06.2025", P["authority_date_ru"])
    return replace_blanks(text, [P["student"]])


def intro_fibs_en(text):
    text = text.replace("Galeeva A.K.", P["vice"])
    text = text.replace("No. 28", f"No. {P['authority_number']}")
    text = text.replace("06/11/2025", P["authority_date_en"])
    return replace_blanks(text, [P["student"]])


def apply_edits(doc, edits):
    paragraphs = list(iter_paragraphs(doc))
    for number, edit in edits.items():
        paragraph = paragraphs[number - 1]
        set_paragraph_text(paragraph, edit(paragraph.text))


def two_party_edits(*, kk_total, ru_total, kk_years, ru_years, kk_sign, ru_sign):
    edits = {
        3: number_line,
        7: lambda text: date_line(text, P["date_kk"]),
        9: intro_two_party_kk,
        16: lambda text: faculty_line(text, P["faculty_kk"]),
        17: lambda text: program_line(text, P["program_kk"]),
        kk_total: lambda text: amount_line(text, P["tuition_kk"]),
        kk_sign[0]: lambda _text: P["vice"],
        kk_sign[1]: student_line,
        kk_sign[2]: address_line,
        kk_sign[3]: clear,
        kk_sign[4]: phone_line,
        kk_sign[5]: email_line,
        kk_sign[6]: student_line,
        kk_sign[7]: consent_line,
        kk_sign[8]: lambda _text: P["date_kk"],
        kk_sign[9]: receipt_kk,
        kk_sign[10]: number_line,
        kk_sign[11]: lambda text: date_line(text, P["date_ru"]),
        kk_sign[11] + 2: intro_two_party_ru,
        kk_sign[12]: lambda text: program_line(text, P["program_ru"]),
        kk_sign[13]: clear,
        kk_sign[14]: lambda text: faculty_line(text, P["faculty_ru"]),
        ru_total: lambda text: amount_line(text, P["tuition_ru"]),
        ru_total + 1: clear,
        ru_sign[0]: lambda _text: P["vice"],
        ru_sign[1]: student_line,
        ru_sign[2]: address_line,
        ru_sign[3]: phone_line,
        ru_sign[4]: email_line,
        ru_sign[5]: student_line,
        ru_sign[6]: consent_line,
        ru_sign[7]: lambda _text: P["date_ru"],
        ru_sign[8]: receipt_ru,
    }

    for paragraph_number, placeholder in kk_years:
        edits[paragraph_number] = lambda text, placeholder=placeholder: amount_line(text, placeholder)
    for paragraph_number, placeholder in ru_years:
        edits[paragraph_number] = lambda text, placeholder=placeholder: amount_line(text, placeholder)

    return edits


def tripartite_edits(
    *,
    kk_total,
    ru_total,
    kk_years=(),
    ru_years=(),
    kk_sign,
    ru_sign,
):
    edits = {
        3: number_line,
        6: lambda text: date_line(text, P["date_kk"]),
        8: intro_tripartite_kk,
        11: student_line,
        13: customer_line,
        16: customer_representative_line,
        17: student_line,
        24: lambda text: faculty_line(text, P["faculty_kk"]),
        25: lambda text: program_line(text, P["program_kk"]),
        kk_total: lambda text: amount_line(text, P["tuition_kk"]),
        kk_sign[0]: lambda _text: P["vice"],
        kk_sign[1]: student_line,
        kk_sign[2]: address_line,
        kk_sign[3]: clear,
        kk_sign[4]: phone_line,
        kk_sign[5]: email_line,
        kk_sign[6]: student_line,
        kk_sign[7]: customer_address_line,
        kk_sign[8]: clear,
        kk_sign[9]: customer_email_line,
        kk_sign[10]: customer_line,
        kk_sign[11]: consent_line,
        kk_sign[12]: lambda _text: P["date_kk"],
        kk_sign[13]: number_line,
        kk_sign[14]: lambda text: date_line(text, P["date_ru"]),
        kk_sign[14] + 2: intro_tripartite_ru,
        kk_sign[15]: student_line,
        kk_sign[16]: clear,
        kk_sign[17]: customer_line,
        kk_sign[18]: customer_representative_line,
        kk_sign[19]: customer_authority_student_line,
        kk_sign[20]: clear,
        kk_sign[21]: lambda text: program_line(text, P["program_ru"]),
        kk_sign[22]: clear,
        kk_sign[23]: lambda text: faculty_line(text, P["faculty_ru"]),
        ru_total: lambda text: amount_line(text, P["tuition_ru"]),
        ru_total + 1: clear,
        ru_sign[0]: lambda _text: P["vice"],
        ru_sign[1]: student_line,
        ru_sign[2]: address_line,
        ru_sign[3]: phone_line,
        ru_sign[4]: email_line,
        ru_sign[5]: student_line,
        ru_sign[6]: customer_address_line,
        ru_sign[7]: customer_email_line,
        ru_sign[8]: customer_line,
        ru_sign[9]: consent_line,
        ru_sign[10]: lambda _text: P["date_ru"],
    }

    for paragraph_number, placeholder in kk_years:
        edits[paragraph_number] = lambda text, placeholder=placeholder: amount_line(text, placeholder)
    for paragraph_number, placeholder in ru_years:
        edits[paragraph_number] = lambda text, placeholder=placeholder: amount_line(text, placeholder)

    return edits


def prepare_fibs(doc):
    apply_edits(
        doc,
        {
            6: number_line,
            12: number_line,
            15: lambda text: date_line(text, P["date_ru"]),
            19: lambda text: date_line(text, P["date_en"]),
            23: intro_fibs_ru,
            24: lambda _text: f"{P['student']}, ИИН {P['student_iin']}",
            30: intro_fibs_en,
            31: lambda _text: f"{P['student']}, IIN {P['student_iin']}",
            36: lambda text: program_line(text, P["program_ru"]),
            40: lambda text: amount_line(text, P["duration_ru"]),
            42: lambda text: program_line(text, P["program_en"]),
            45: lambda text: amount_line(text, P["duration_en"]),
            60: lambda text: amount_line(text, P["qualification_ru"]),
            95: lambda text: amount_line(text, P["qualification_en"]),
            176: lambda text: amount_line(text, P["tuition_ru"]),
            189: lambda text: amount_line(text, P["tuition_en"]),
            395: lambda _text: f"Член Правления-Проректор по академическим вопросам {P['vice']} ______________________________________(подпись)",
            399: lambda _text: f"{P['student']}, {P['date_ru']}",
            404: lambda text: replace_blanks(text, [P["student_citizenship"]]),
            407: lambda _text: f"{P['student_identity']}",
            408: lambda _text: f"{P['student_address']}, {P['student_phone']}, {P['student_email']}",
            412: lambda _text: f"{P['student']} ______________________________________(signature)",
            416: lambda _text: f"{P['student']}, {P['date_en']}",
            417: lambda _text: f"{P['student_address']}, {P['student_phone']}, {P['student_email']}",
            422: lambda text: replace_blanks(text, [P["student_citizenship"]]),
        },
    )


def prepare_document(path, output_path):
    doc = Document(path)
    name = path.name

    if "FIBS" in name:
        prepare_fibs(doc)
    elif "трехсторонний" in name:
        if "EMBA" in name and "50.50" in name:
            edits = tripartite_edits(
                kk_total=148,
                ru_total=342,
                kk_sign=(209, 213, 215, 216, 217, 218, 221, 226, 227, 228, 229, 231, 232, 235, 239, 243, 244, 246, 247, 249, 250, 256, 257, 259),
                ru_sign=(388, 392, 394, 395, 396, 399, 404, 405, 406, 409, 410),
            )
        elif "EMBA" in name:
            edits = tripartite_edits(
                kk_total=143,
                ru_total=338,
                kk_sign=(205, 209, 211, 212, 213, 214, 217, 222, 223, 224, 225, 227, 228, 231, 235, 239, 240, 242, 243, 245, 246, 252, 253, 255),
                ru_sign=(384, 388, 390, 391, 392, 395, 400, 401, 402, 405, 406),
            )
        elif "DBA" in name and "50.50" in name:
            edits = tripartite_edits(
                kk_total=145,
                ru_total=351,
                kk_years=((151, P["year_1_kk"]), (153, P["year_2_kk"]), (155, P["year_3_kk"])),
                ru_years=((356, P["year_1_ru"]), (357, P["year_2_ru"]), (358, P["year_3_ru"])),
                kk_sign=(215, 219, 221, 222, 223, 224, 227, 232, 233, 234, 235, 237, 238, 240, 244, 248, 249, 251, 252, 254, 255, 261, 262, 264),
                ru_sign=(400, 404, 406, 407, 408, 411, 416, 417, 418, 420, 421),
            )
        elif "DBA" in name:
            edits = tripartite_edits(
                kk_total=142,
                ru_total=342,
                kk_years=((148, P["year_1_kk"]), (150, P["year_2_kk"]), (152, P["year_3_kk"])),
                ru_years=((347, P["year_1_ru"]), (348, P["year_2_ru"]), (349, P["year_3_ru"])),
                kk_sign=(208, 212, 214, 215, 216, 217, 220, 225, 226, 227, 228, 230, 231, 233, 237, 241, 242, 244, 245, 247, 248, 254, 255, 257),
                ru_sign=(392, 396, 398, 399, 400, 403, 408, 409, 410, 412, 413),
            )
        else:
            edits = tripartite_edits(
                kk_total=140,
                ru_total=344,
                kk_years=((146, P["year_1_kk"]), (148, P["year_2_kk"])),
                ru_years=((349, P["year_1_ru"]), (350, P["year_2_ru"])),
                kk_sign=(208, 212, 214, 215, 216, 217, 220, 225, 226, 227, 228, 230, 231, 233, 237, 241, 242, 244, 245, 247, 248, 254, 255, 257),
                ru_sign=(392, 396, 398, 399, 400, 403, 408, 409, 410, 412, 413),
            )
        apply_edits(doc, edits)
    else:
        if "EMBA" in name:
            edits = two_party_edits(
                kk_total=114,
                ru_total=265,
                kk_years=(),
                ru_years=((271, P["year_1_ru"]),),
                kk_sign=(167, 170, 172, 173, 174, 175, 177, 179, 180, 182, 184, 189, 197, 198, 200),
                ru_sign=(315, 318, 320, 321, 322, 324, 326, 327, 330),
            )
            edits.update({
                119: installment_line(1, 1),
                120: installment_line(1, 2),
                121: installment_line(1, 3),
                122: installment_line(1, 4),
                123: installment_line(1, 5),
                124: installment_line(1, 6),
            })
        elif "МBA" in name:
            edits = two_party_edits(
                kk_total=107,
                ru_total=290,
                kk_years=((112, P["year_1_kk"]), (120, P["year_2_kk"])),
                ru_years=((295, P["year_1_ru"]), (303, P["year_2_ru"])),
                kk_sign=(188, 191, 193, 194, 195, 196, 198, 200, 201, 203, 205, 209, 217, 218, 220),
                ru_sign=(353, 356, 358, 359, 360, 362, 364, 365, 368),
            )
            edits.update({
                113: installment_line(1, 1),
                114: installment_line(1, 2),
                115: installment_line(1, 3),
                116: installment_line(1, 4),
                117: installment_line(1, 5),
                118: installment_line(1, 6),
                121: installment_line(2, 1),
                122: installment_line(2, 2),
                123: installment_line(2, 3),
                124: installment_line(2, 4),
                125: installment_line(2, 5),
                126: installment_line(2, 6),
            })
        else:
            edits = two_party_edits(
                kk_total=91,
                ru_total=284,
                kk_years=((96, P["year_1_kk"]), (103, P["year_2_kk"]), (111, P["year_3_kk"])),
                ru_years=((289, P["year_1_ru"]), (297, P["year_2_ru"]), (305, P["year_3_ru"])),
                kk_sign=(182, 185, 187, 188, 189, 190, 192, 194, 195, 197, 199, 203, 211, 212, 214),
                ru_sign=(353, 356, 358, 359, 360, 362, 364, 365, 368),
            )
            edits.update({
                97: installment_line(1, 1),
                98: installment_line(1, 2),
                99: installment_line(1, 3),
                100: installment_line(1, 4),
                101: installment_line(1, 5),
                104: installment_line(2, 1),
                105: installment_line(2, 2),
                106: installment_line(2, 3),
                107: installment_line(2, 4),
                108: installment_line(2, 5),
                109: installment_line(2, 6),
                112: installment_line(3, 1),
                113: installment_line(3, 2),
                114: installment_line(3, 3),
                115: installment_line(3, 4),
                116: installment_line(3, 5),
                117: installment_line(3, 6),
            })
        apply_edits(doc, edits)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for path in sorted(SOURCE_DIR.glob("*.docx")):
        output_path = OUTPUT_DIR / f"{path.stem}_prepared.docx"
        prepare_document(path, output_path)
        print(output_path)


if __name__ == "__main__":
    main()
