from pathlib import Path

from docx import Document


BASE = Path("prepared_templates/trustme_docs")
VICE = "{{ university_representative_full_name }}"


def iter_paragraphs(parent):
    yield from parent.paragraphs
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def set_text(paragraph, text):
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return

    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def ru_tranche(year, tranche, label):
    return (
        f"{label} транш в размере {{{{ year_{year}_tranche_{tranche}_amount }}}}тг. "
        f"до {{{{ year_{year}_tranche_{tranche}_due_date }}}}г."
    )


def main():
    paths = sorted(BASE.glob("*.docx"), key=lambda path: path.name)

    # DBA two-party: the reviewed file had one missed due-date placeholder.
    dba = Document(paths[0])
    dba_paragraphs = list(iter_paragraphs(dba))
    set_text(dba_paragraphs[293 - 1], ru_tranche(1, 4, "4. Четвертый"))
    dba.save(paths[0])

    # EMBA two-party: complete Russian payment schedule, preserving the Kazakh
    # explanatory sentence that shares a paragraph with the first tranche.
    emba = Document(paths[3])
    emba_paragraphs = list(iter_paragraphs(emba))
    emba_replacements = {
        119: (
            "3.3.1. EMBA — оқу басталғанға дейін күнтізбелік 10 күннен кешіктірмей, "
            "бекітілген академиялық күнтізбеге сәйкес жалпы соманың 20%-ы төленеді, "
            "ал қалған сома 6 ай ішінде тең үлестермен төленеді. "
            "1. Бірінші транш — {{ year_1_tranche_1_amount }} теңге мөлшерінде, "
            "{{ year_1_tranche_1_due_date }} дейін."
        ),
        271: ru_tranche(1, 1, "1.Первый"),
        273: ru_tranche(1, 2, "2. Второй"),
        274: ru_tranche(1, 3, "3. Третий"),
        275: ru_tranche(1, 4, "4. Четвертый"),
        276: ru_tranche(1, 5, "5. Пятый"),
        277: ru_tranche(1, 6, "6. Шестой"),
        314: f"Член Правления – проректор {VICE}",
        315: "",
    }
    for number, value in emba_replacements.items():
        set_text(emba_paragraphs[number - 1], value)
    emba.save(paths[3])

    # FIBS: repair variables accidentally split by a placeholder cleanup pass.
    fibs = Document(paths[6])
    fibs_paragraphs = list(iter_paragraphs(fibs))
    set_text(
        fibs_paragraphs[395 - 1],
        f"Член Правления-Проректор по академическим вопросам {VICE} (подпись)",
    )
    set_text(fibs_paragraphs[411 - 1], VICE)
    set_text(fibs_paragraphs[412 - 1], "{{ side_1_full_name }} (signature)")
    fibs.save(paths[6])

    # MBA two-party: complete Russian payment schedule.
    mba = Document(paths[7])
    mba_paragraphs = list(iter_paragraphs(mba))
    mba_replacements = {
        296: ru_tranche(1, 1, "1.Первый"),
        297: ru_tranche(1, 2, "2. Второй"),
        298: ru_tranche(1, 3, "3. Третий"),
        299: ru_tranche(1, 4, "4. Четвертый"),
        300: ru_tranche(1, 5, "5. Пятый"),
        301: ru_tranche(1, 6, "6. Шестой"),
        304: ru_tranche(2, 1, "1.Первый"),
        305: ru_tranche(2, 2, "2. Второй"),
        306: ru_tranche(2, 3, "3. Третий"),
        307: ru_tranche(2, 4, "4. Четвертый"),
        308: ru_tranche(2, 5, "5. Пятый"),
        309: ru_tranche(2, 6, "6. Шестой"),
    }
    for number, value in mba_replacements.items():
        set_text(mba_paragraphs[number - 1], value)
    mba.save(paths[7])

    print("TrustMe prepared templates fixed.")


if __name__ == "__main__":
    main()
