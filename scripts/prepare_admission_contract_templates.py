from pathlib import Path
import re

from docx import Document


P = {
    "contract_number": "{{ contract_number }}",
    "date_kk": "{{ contract_date_text_kk }}",
    "date_ru": "{{ contract_date_text_ru }}",
    "student": "{{ side_1_full_name }}",
    "student_iin": "{{ side_1_iin_bin }}",
    "student_address": "{{ student_address }}",
    "student_phone": "{{ side_1_phone }}",
    "student_email": "{{ side_1_email }}",
    "parent": "{{ student_parent_full_name }}",
    "parent_details_ru": "{{ student_parent_details_ru }}",
    "parent_details_kk": "{{ student_parent_details_kk }}",
    "program_kk": "{{ program_code }} {{ program_name_kk }}",
    "program_ru": "{{ program_code }} {{ program_name_ru }}",
    "faculty_kk": "{{ program_faculty_kk }}",
    "faculty_ru": "{{ program_faculty_ru }}",
    "grant_number": "{{ grant_number }}",
    "vice": "{{ university_representative_full_name }}",
    "tuition_kk": "{{ tuition_amount_full_kk }}",
    "tuition_ru": "{{ tuition_amount_full_ru }}",
    "parent_iin": "{{ student_parent_iin }}",
    "parent_phone": "{{ student_parent_phone }}",
    "parent_address": "{{ student_parent_address }}",
}


def iter_all_paragraphs(doc):
    yield from doc.paragraphs

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def nonempty_paragraphs(doc):
    return [paragraph for paragraph in iter_all_paragraphs(doc) if paragraph.text.strip()]


def set_paragraph_text(paragraph, text):
    text = tidy(text)
    text = normalize_kazakh_letters(text)
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return

    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def tidy(text):
    text = re.sub(r"(?<![\s\u00ab]){{", r" {{", text)
    text = re.sub(r"}}(?![\s,.;:\u00bb)])", r"}} ", text)
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    return text.strip()


def normalize_kazakh_letters(text):
    return text.replace("Ə", "Ә").replace("ə", "ә")


def normalize_document_text(doc):
    for paragraph in iter_all_paragraphs(doc):
        for run in paragraph.runs:
            normalized = normalize_kazakh_letters(run.text)
            if normalized != run.text:
                run.text = normalized


def replace_blanks(text, replacements):
    parts = re.split(r"(_{2,})", text)
    result = []
    replacement_index = 0

    for part in parts:
        if re.fullmatch(r"_{2,}", part):
            if replacement_index < len(replacements):
                result.append(replacements[replacement_index])
                replacement_index += 1
        else:
            result.append(part)

    return "".join(result)


def contract_date_line(text, placeholder):
    prefix = text.split("\u00ab", 1)[0].strip()
    return f"{prefix} {placeholder}" if prefix else placeholder


def edit_number(text):
    return replace_blanks(text, [P["contract_number"]])


def edit_student(text):
    return replace_blanks(text, [P["student"]])


def edit_student_requisites_kk(text):
    return replace_blanks(text, [P["student"]])


def edit_student_requisites_ru(text):
    return replace_blanks(text, [P["student"]])


def edit_address(text):
    return replace_blanks(text, [P["student_address"]])


def edit_phone(text):
    return replace_blanks(text, [P["student_phone"]])


def edit_email(text):
    return replace_blanks(text, [P["student_email"]])


def edit_faculty(text, placeholder):
    return replace_blanks(text, [placeholder])


def edit_consent(text):
    return replace_blanks(text, [P["student"]])


def edit_parent_and_child(text):
    return replace_blanks(text, [parent_details_ru(), P["student"]])


def edit_student_intro(text):
    return replace_blanks(text, [P["student"]]).rstrip(" ,") + ","


def grant_program_kk():
    return f"{P['program_kk']}, білім гранты куәлігі № {P['grant_number']}"


def grant_program_ru():
    return f"{P['program_ru']}, свидетельство о гранте № {P['grant_number']}"


def parent_details_ru():
    return P["parent_details_ru"]


def parent_details_kk():
    return P["parent_details_kk"]


def receipt_ru():
    return f"Экземпляр Договора получил(-а) {P['student']} / {P['date_ru']} /"


def apply_edits(doc, edits):
    paragraphs = nonempty_paragraphs(doc)
    for number, edit in edits.items():
        paragraph = paragraphs[number - 1]
        set_paragraph_text(paragraph, edit(paragraph.text))


def prepare_bachelor_grant(doc):
    apply_edits(doc, {
        2: edit_number,
        4: lambda text: contract_date_line(text, P["date_kk"]),
        6: edit_student,
        11: lambda text: edit_faculty(text, P["faculty_kk"]),
        12: lambda text: grant_program_kk(),
        117: lambda text: P["vice"],
        119: edit_student_requisites_kk,
        121: edit_address,
        122: edit_phone,
        123: edit_email,
        124: edit_student,
        126: lambda text: P["date_kk"],
        127: edit_consent,
        128: lambda text: replace_blanks(text, [parent_details_kk(), P["student"]]),
        131: edit_number,
        134: lambda text: contract_date_line(text, P["date_ru"]),
        136: edit_student_intro,
        141: lambda text: grant_program_ru(),
        143: lambda text: edit_faculty(text, P["faculty_ru"]),
        247: lambda text: P["vice"],
        249: edit_student_requisites_ru,
        251: edit_address,
        252: edit_phone,
        253: edit_email,
        254: edit_student,
        256: lambda text: P["date_ru"],
        257: edit_consent,
        258: lambda text: replace_blanks(text, [parent_details_ru(), P["student"]]),
        259: lambda text: receipt_ru(),
    })


def prepare_paid(doc):
    apply_edits(doc, {
        2: edit_number,
        3: lambda text: contract_date_line(text, P["date_kk"]),
        5: edit_student,
        10: lambda text: edit_faculty(text, P["faculty_kk"]),
        11: lambda text: P["program_kk"],
        81: lambda text: P["tuition_kk"],
        133: lambda text: P["vice"],
        135: edit_student_requisites_kk,
        137: edit_address,
        138: edit_phone,
        139: edit_email,
        140: edit_student,
        142: lambda text: P["date_kk"],
        143: edit_consent,
        144: lambda text: replace_blanks(text, [parent_details_kk()]),
        145: lambda text: replace_blanks(text, [P["student"]]),
        147: edit_number,
        149: lambda text: contract_date_line(text, P["date_ru"]),
        151: edit_student_intro,
        156: lambda text: P["program_ru"],
        158: lambda text: edit_faculty(text, P["faculty_ru"]),
        227: lambda text: P["tuition_ru"],
        279: lambda text: P["vice"],
        281: edit_student_requisites_ru,
        283: edit_address,
        284: edit_phone,
        285: edit_email,
        286: edit_student,
        288: lambda text: P["date_ru"],
        289: edit_consent,
        290: lambda text: replace_blanks(text, [parent_details_ru()]),
        291: lambda text: replace_blanks(text, [P["student"]]),
        292: lambda text: receipt_ru(),
    })


def prepare_doctoral_grant(doc):
    apply_edits(doc, {
        2: edit_number,
        4: lambda text: contract_date_line(text, P["date_kk"]),
        6: edit_student,
        10: lambda text: edit_faculty(text, P["faculty_kk"]),
        11: lambda text: grant_program_kk(),
        109: lambda text: P["vice"],
        111: edit_student_requisites_kk,
        113: edit_address,
        114: edit_phone,
        115: edit_email,
        116: edit_student,
        118: edit_consent,
        119: lambda text: P["date_kk"],
        122: edit_number,
        125: lambda text: contract_date_line(text, P["date_ru"]),
        127: edit_student_intro,
        132: lambda text: grant_program_ru(),
        134: lambda text: edit_faculty(text, P["faculty_ru"]),
        232: lambda text: P["vice"],
        234: edit_student_requisites_ru,
        236: edit_address,
        237: edit_phone,
        238: edit_email,
        239: edit_student,
        241: edit_consent,
        242: lambda text: P["date_ru"],
        243: lambda text: receipt_ru(),
    })


def prepare_master_grant(doc):
    apply_edits(doc, {
        2: edit_number,
        4: lambda text: contract_date_line(text, P["date_kk"]),
        6: edit_student,
        10: lambda text: edit_faculty(text, P["faculty_kk"]),
        11: lambda text: grant_program_kk(),
        111: lambda text: P["vice"],
        113: edit_student_requisites_kk,
        115: edit_address,
        116: edit_phone,
        117: edit_email,
        118: edit_student,
        120: edit_consent,
        121: lambda text: P["date_kk"],
        123: edit_number,
        126: lambda text: contract_date_line(text, P["date_ru"]),
        128: edit_student_intro,
        132: lambda text: replace_blanks(text, [grant_program_ru()]),
        134: lambda text: edit_faculty(text, P["faculty_ru"]),
        234: lambda text: P["vice"],
        236: edit_student_requisites_ru,
        238: edit_address,
        239: edit_phone,
        240: edit_email,
        241: edit_student,
        243: edit_consent,
        244: lambda text: P["date_ru"],
        245: lambda text: receipt_ru(),
    })


def find_docx(directory, marker, *, prepared):
    for path in directory.glob("*.docx"):
        if marker in path.name and ("prepared" in path.name) == prepared:
            return path
    raise FileNotFoundError(f"DOCX with marker {marker!r} was not found in {directory}")


def prepare(marker, prepare_func):
    source_path = find_docx(Path.home() / "Downloads", marker, prepared=False)
    output_path = find_docx(Path("prepared_templates") / "admissions", marker, prepared=True)

    doc = Document(source_path)
    prepare_func(doc)
    normalize_document_text(doc)
    doc.save(output_path)
    return output_path


def validate(path):
    doc = Document(path)
    text = "\n".join(paragraph.text for paragraph in iter_all_paragraphs(doc))
    variables = sorted(set(re.findall(r"{{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*}}", text)))
    return text.count("?"), variables


def main():
    jobs = [
        ("бак_грант", prepare_bachelor_grant),
        ("всеуровни", prepare_paid),
        ("док_грант", prepare_doctoral_grant),
        ("маг_грант", prepare_master_grant),
    ]
    for marker, prepare_func in jobs:
        output = prepare(marker, prepare_func)
        question_marks, variables = validate(output)
        print(str(output).encode("unicode_escape").decode("ascii"))
        print(f"question_marks={question_marks}")
        print(f"variables={len(variables)}")
        print(", ".join(variables))


if __name__ == "__main__":
    main()
