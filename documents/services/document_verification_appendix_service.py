import tempfile
from io import BytesIO
from pathlib import Path
from urllib.parse import urljoin

import qrcode
from django.conf import settings
from django.urls import reverse
from docx import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


class DocumentVerificationAppendixService:
    MARKER = "DOCUMENT_VERIFICATION_PAGE_START"
    LEGACY_MARKERS = ("QOLQOYU_VERIFICATION_PAGE_START",)

    @classmethod
    def finalize_docx(cls, *, document, file_path: str, request=None) -> None:
        from documents.services.docx_template_service import DocxTemplateService

        DocxTemplateService.add_page_numbering(file_path)
        cls.append_verification_page(
            document=document,
            file_path=file_path,
            request=request,
        )

    @classmethod
    def finalize_docx_bytes(cls, *, document, content: bytes, request=None) -> bytes:
        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "document.docx"
            file_path.write_bytes(content)

            cls.finalize_docx(
                document=document,
                file_path=str(file_path),
                request=request,
            )

            return file_path.read_bytes()

    @classmethod
    def append_verification_page(cls, *, document, file_path: str, request=None) -> None:
        doc = DocxDocument(file_path)
        cls.remove_existing_verification_page(doc)

        verification_url = cls.build_verification_url(document=document, request=request)

        marker_paragraph = doc.add_paragraph()
        marker_paragraph.add_run().add_break(WD_BREAK.PAGE)
        marker_run = marker_paragraph.add_run(cls.MARKER)
        marker_run.font.hidden = True

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("Проверка документа / Құжатты тексеру")
        title_run.bold = True
        title_run.font.size = Pt(16)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Document verification")
        subtitle_run.font.size = Pt(10)

        cls.add_metadata_table(
            doc=doc,
            document=document,
            verification_url=verification_url,
        )
        cls.add_qr_notice(
            doc=doc,
            verification_url=verification_url,
        )

        doc.save(file_path)

    @classmethod
    def remove_existing_verification_page(cls, doc) -> None:
        body = doc._element.body
        children = list(body)
        marker_index = None

        for index, child in enumerate(children):
            child_text = cls.element_text(child)
            if cls.MARKER in child_text or any(marker in child_text for marker in cls.LEGACY_MARKERS):
                marker_index = index
                break

        if marker_index is None:
            return

        for child in children[marker_index:]:
            if child.tag == qn("w:sectPr"):
                continue
            body.remove(child)

    @staticmethod
    def element_text(element) -> str:
        return "".join(text_node.text or "" for text_node in element.iter(qn("w:t")))

    @classmethod
    def add_metadata_table(cls, *, doc, document, verification_url: str) -> None:
        table = doc.add_table(rows=0, cols=2)
        cls.apply_table_grid_style(doc=doc, table=table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        external_reference = cls.get_external_reference(document)
        rows = [
            ("Тип документа", "Договор / Contract"),
            ("Название документа", document.title or "-"),
            ("Внешний ID", external_reference or "-"),
            ("Номер договора", document.contract_number or "-"),
            ("Дата договора", document.get_contract_date_display()),
            ("Статус", document.get_status_display()),
            ("Страница проверки", verification_url),
        ]

        for label, value in rows:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value

            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        doc.add_paragraph()

    @classmethod
    def add_qr_notice(cls, *, doc, verification_url: str) -> None:
        table = doc.add_table(rows=1, cols=2)
        cls.apply_table_grid_style(doc=doc, table=table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        qr_cell = table.cell(0, 0)
        qr_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        qr_paragraph = qr_cell.paragraphs[0]
        qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qr_paragraph.add_run().add_picture(
            cls.build_qr_image(verification_url),
            width=Inches(1.45),
        )

        notice_cell = table.cell(0, 1)
        notice_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        notice_paragraph = notice_cell.paragraphs[0]
        notice_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        notice_text = (
            "Отсканируйте QR-код для проверки подлинности и целостности документа. "
            "QR-код ведет на публичную страницу проверки документа, где отображаются "
            "номер договора, статус документа, сохраненный хэш и запись ledger при наличии."
        )
        run = notice_paragraph.add_run(notice_text)
        run.font.size = Pt(9)

        url_paragraph = notice_cell.add_paragraph()
        url_run = url_paragraph.add_run(verification_url)
        url_run.font.size = Pt(8)

    @staticmethod
    def apply_table_grid_style(*, doc, table) -> None:
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.TABLE and style.name == "Table Grid":
                table.style = style
                return

    @classmethod
    def get_external_reference(cls, document) -> str:
        for related_name in ["admission_contracts", "admission_applications"]:
            manager = getattr(document, related_name, None)

            if manager is None:
                continue

            try:
                contract = manager.order_by("-created_at").first()
            except Exception:
                continue

            if contract and contract.external_id:
                return contract.external_id

        return ""

    @staticmethod
    def build_qr_image(data: str) -> BytesIO:
        qr = qrcode.QRCode(
            version=None,
            error_correction=qrcode.constants.ERROR_CORRECT_M,
            box_size=8,
            border=2,
        )
        qr.add_data(data)
        qr.make(fit=True)
        image = qr.make_image(fill_color="black", back_color="white").convert("RGB")

        output = BytesIO()
        image.save(output, format="PNG")
        output.seek(0)
        return output

    @classmethod
    def build_verification_url(cls, *, document, request=None) -> str:
        path = reverse("documents:document_verification", args=[document.verification_token])

        if request is not None:
            return request.build_absolute_uri(path)

        public_site_url = getattr(settings, "PUBLIC_SITE_URL", "").strip()

        if public_site_url:
            return urljoin(public_site_url.rstrip("/") + "/", path.lstrip("/"))

        for origin in getattr(settings, "CSRF_TRUSTED_ORIGINS", []):
            origin = origin.strip()

            if not origin:
                continue

            if any(local_host in origin for local_host in ["localhost", "127.0.0.1", "host.docker.internal"]):
                continue

            return urljoin(origin.rstrip("/") + "/", path.lstrip("/"))

        django_url = getattr(settings, "ONLYOFFICE_DJANGO_URL", "").strip()

        if django_url:
            return urljoin(django_url.rstrip("/") + "/", path.lstrip("/"))

        return path
