from bs4 import BeautifulSoup
from docx import Document as DocxDocument

import os
import sys
import tempfile
from pathlib import Path

from .template_file_service import TemplateFileService


class HtmlToDocxService:
    @classmethod
    def render_html_to_docx(cls, html: str, output_path: str):
        if cls.render_html_to_docx_with_libreoffice(html=html, output_path=output_path):
            return

        soup = BeautifulSoup(html or "", "html.parser")
        doc = DocxDocument()

        body_elements = soup.find_all(["h1", "h2", "h3", "p", "div", "table"], recursive=True)

        for element in body_elements:
            if element.name == "h1":
                doc.add_heading(element.get_text(" ", strip=True), level=1)

            elif element.name == "h2":
                doc.add_heading(element.get_text(" ", strip=True), level=2)

            elif element.name == "h3":
                doc.add_heading(element.get_text(" ", strip=True), level=3)

            elif element.name in ["p", "div"]:
                text = element.get_text(" ", strip=True)
                if text:
                    doc.add_paragraph(text)

            elif element.name == "table":
                rows = element.find_all("tr")
                if not rows:
                    continue

                first_row_cells = rows[0].find_all(["td", "th"])
                if not first_row_cells:
                    continue

                table = doc.add_table(
                    rows=len(rows),
                    cols=len(first_row_cells)
                )

                table.style = "Table Grid"

                for row_index, row in enumerate(rows):
                    cells = row.find_all(["td", "th"])

                    for cell_index, cell in enumerate(cells):
                        if cell_index < len(table.rows[row_index].cells):
                            table.rows[row_index].cells[cell_index].text = cell.get_text(" ", strip=True)

        doc.save(output_path)

    @classmethod
    def render_html_to_docx_with_libreoffice(cls, *, html: str, output_path: str) -> bool:
        if "test" in sys.argv:
            return False

        soffice_path = TemplateFileService.find_soffice()

        if not soffice_path:
            return False

        with tempfile.TemporaryDirectory() as temp_dir:
            html_path = os.path.join(temp_dir, "document.html")
            converted_path = os.path.join(temp_dir, "document.docx")

            with open(html_path, "w", encoding="utf-8") as html_file:
                html_file.write(html or "")

            result = TemplateFileService.run_soffice_conversion(
                soffice_path=soffice_path,
                file_path=html_path,
                output_dir=temp_dir,
                convert_to="docx",
                timeout=60,
                operation="converting HTML to DOCX",
            )

            if result.returncode != 0 or not os.path.exists(converted_path):
                alternate_path = os.path.join(temp_dir, Path(html_path).stem + ".docx")

                if not os.path.exists(alternate_path):
                    return False

                converted_path = alternate_path

            with open(converted_path, "rb") as converted_file:
                with open(output_path, "wb") as output_file:
                    output_file.write(converted_file.read())

        return True
