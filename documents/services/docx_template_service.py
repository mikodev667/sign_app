import re
from zipfile import ZIP_DEFLATED, ZipFile
from xml.sax.saxutils import escape as xml_escape

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxtpl import DocxTemplate


VARIABLE_PATTERN = re.compile(r"{{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*}}")


class DocxTemplateService:
    @classmethod
    def extract_text_from_docx(cls, file_path: str) -> str:
        doc = DocxDocument(file_path)
        parts = []

        for paragraph in doc.paragraphs:
            parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)

        return "\n".join(parts)

    @classmethod
    def extract_variables(cls, file_path: str) -> list[str]:
        text = cls.extract_text_from_docx(file_path)
        variables = VARIABLE_PATTERN.findall(text)

        result = []
        seen = set()

        for variable in variables:
            if variable not in seen:
                seen.add(variable)
                result.append(variable)

        return result

    @classmethod
    def render_docx(cls, *, template_path: str, output_path: str, values: dict):
        if cls.can_render_with_simple_replacements(template_path):
            cls.render_docx_with_simple_replacements(
                template_path=template_path,
                output_path=output_path,
                values=values,
            )
        else:
            doc = DocxTemplate(template_path)
            doc.render(values)
            doc.save(output_path)

        cls.normalize_rendered_value_styles(output_path, values)
        cls.add_page_numbering(output_path)

    @classmethod
    def can_render_with_simple_replacements(cls, template_path: str) -> bool:
        xml_text = cls.read_docx_word_xml(template_path)

        if "{%" in xml_text or "{#" in xml_text:
            return False

        variables = cls.extract_variables(template_path)
        return all(
            re.search(cls.variable_xml_pattern(variable), xml_text)
            for variable in variables
        )

    @classmethod
    def render_docx_with_simple_replacements(cls, *, template_path: str, output_path: str, values: dict):
        safe_values = {
            key: xml_escape(str(value or ""))
            for key, value in (values or {}).items()
        }

        with ZipFile(template_path, "r") as source, ZipFile(output_path, "w", ZIP_DEFLATED) as target:
            for item in source.infolist():
                data = source.read(item.filename)

                if cls.is_word_xml_file(item.filename):
                    text = data.decode("utf-8")

                    for variable in set(VARIABLE_PATTERN.findall(text)):
                        replacement = safe_values.get(variable, "")
                        text = re.sub(
                            cls.variable_xml_pattern(variable),
                            lambda _match, replacement=replacement: replacement,
                            text,
                        )

                    data = text.encode("utf-8")

                target.writestr(item, data)

    @classmethod
    def read_docx_word_xml(cls, file_path: str) -> str:
        parts = []

        with ZipFile(file_path, "r") as archive:
            for item in archive.infolist():
                if cls.is_word_xml_file(item.filename):
                    parts.append(archive.read(item.filename).decode("utf-8"))

        return "\n".join(parts)

    @staticmethod
    def is_word_xml_file(file_name: str) -> bool:
        return file_name.startswith("word/") and file_name.endswith(".xml")

    @staticmethod
    def variable_xml_pattern(variable: str) -> str:
        return r"{{\s*" + re.escape(variable) + r"\s*}}"

    @classmethod
    def normalize_rendered_value_styles(cls, file_path: str, values: dict):
        value_texts = [
            str(value).strip()
            for value in (values or {}).values()
            if str(value or "").strip()
        ]

        if not value_texts:
            return

        doc = DocxDocument(file_path)
        changed = False

        for paragraph in cls.iter_all_paragraphs(doc):
            for run in paragraph.runs:
                if not run.text:
                    continue

                if any(value in run.text for value in value_texts):
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.highlight_color = None
                    changed = True

        if changed:
            doc.save(file_path)

    @classmethod
    def add_page_numbering(cls, file_path: str):
        doc = DocxDocument(file_path)
        changed = False

        for section in doc.sections:
            footer = section.footer

            if cls.footer_has_page_numbering(footer):
                continue

            paragraph = footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run("- ")
            cls.add_field(paragraph, "PAGE")
            paragraph.add_run(" / ")
            cls.add_field(paragraph, "NUMPAGES")
            paragraph.add_run(" -")
            changed = True

        if changed:
            doc.save(file_path)

    @classmethod
    def footer_has_page_numbering(cls, footer):
        footer_xml = footer._element.xml
        return "PAGE" in footer_xml and "NUMPAGES" in footer_xml

    @classmethod
    def add_field(cls, paragraph, field_code: str):
        run = paragraph.add_run()

        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        run._r.append(begin)

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = f" {field_code} "
        run._r.append(instr_text)

        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        run._r.append(separate)

        display_text = OxmlElement("w:t")
        display_text.text = "1"
        run._r.append(display_text)

        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(end)

    @classmethod
    def iter_all_paragraphs(cls, doc):
        yield from doc.paragraphs

        for table in doc.tables:
            yield from cls.iter_table_paragraphs(table)

        for section in doc.sections:
            for header_footer in [
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ]:
                yield from header_footer.paragraphs

                for table in header_footer.tables:
                    yield from cls.iter_table_paragraphs(table)

    @classmethod
    def iter_table_paragraphs(cls, table):
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs

                for nested_table in cell.tables:
                    yield from cls.iter_table_paragraphs(nested_table)
