import hashlib
import json
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from types import SimpleNamespace
from unittest.mock import patch

from requests import Response

from django.contrib.auth import get_user_model
from django.core.files.base import ContentFile
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import SimpleTestCase, TestCase, override_settings
from django.urls import reverse
from django.utils import timezone
from docx import Document as WordDocument
from docx.shared import Cm
from docx.shared import RGBColor

from documents.forms import DocumentTemplateUploadForm
from documents.models import (
    Document,
    DocumentFieldValue,
    DocumentLedgerRecord,
    DocumentTemplate,
    StoredObject,
    TemplateParty,
    TemplatePartyField,
)
from documents.services.object_storage_service import ObjectStorageService
from documents.services.docx_template_service import DocxTemplateService
from documents.services.docx_preview_service import DocxPreviewService
from documents.services.document_docx_render_service import DocumentDocxRenderService
from documents.services.document_verification_appendix_service import DocumentVerificationAppendixService
from documents.services.lawvision_service import LawVisionError, LawVisionService
from documents.services.money_amount_service import MoneyAmountService
from documents.services.onlyoffice_service import OnlyOfficeService
from documents.services.document_ledger_service import DocumentLedgerError, DocumentLedgerService
from documents.services.pdf_export_service import ExportedPdf
from documents.services.template_file_service import TemplateFileService
from organizations.models import Department, Organization, OrganizationMember
from signing.models import Signer


class LawVisionServiceTests(SimpleTestCase):
    def test_parse_response_exposes_non_json_response_details(self):
        response = Response()
        response.status_code = 502
        response.headers["Content-Type"] = "text/html"
        response._content = b"<html><body>Bad Gateway</body></html>"

        with self.assertRaises(LawVisionError) as context:
            LawVisionService.parse_response(response)

        self.assertEqual(context.exception.error_code, "invalid_json")
        self.assertEqual(context.exception.status_code, 502)
        self.assertIn("HTTP 502", str(context.exception))
        self.assertIn("text/html", str(context.exception))
        self.assertIn("Bad Gateway", str(context.exception))


class TemplateFileServiceTests(SimpleTestCase):
    def test_validate_file_name_allows_doc_and_docx(self):
        for file_name in ["template.doc", "template.docx"]:
            with self.subTest(file_name=file_name):
                TemplateFileService.validate_file_name(file_name)

    def test_validate_file_name_rejects_pdf_and_other_formats(self):
        for file_name in ["template.pdf", "template.txt"]:
            with self.subTest(file_name=file_name):
                with self.assertRaises(ValueError):
                    TemplateFileService.validate_file_name(file_name)

    def test_extract_variables_from_text_keeps_order_and_removes_duplicates(self):
        variables = TemplateFileService.extract_variables_from_text(
            "Hello {{ full_name }}. IIN {{ iin }}. Again {{ full_name }}."
        )

        self.assertEqual(variables, ["full_name", "iin"])

    def test_as_docx_path_yields_docx_file_without_conversion(self):
        with TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "template.docx"
            docx_path.write_bytes(self.create_docx_bytes())

            with TemplateFileService.as_docx_path(str(docx_path)) as resolved_path:
                self.assertEqual(resolved_path, str(docx_path))

    def test_convert_to_html_reads_docx_content(self):
        with TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "template.docx"
            docx_path.write_bytes(self.create_docx_bytes("Hello {{ full_name }}"))

            with patch.object(
                DocxPreviewService,
                "convert_docx_to_html_with_libreoffice",
                return_value="",
            ):
                html = TemplateFileService.convert_to_html(str(docx_path))

        self.assertIn("Hello {{ full_name }}", html)

    def test_get_page_layout_reads_docx_page_margins(self):
        with TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "template.docx"
            docx_path.write_bytes(self.create_docx_bytes_with_margins())

            layout = DocxPreviewService.get_page_layout(str(docx_path))

        self.assertEqual(layout["margin_top"], "1")
        self.assertEqual(layout["margin_right"], "1.5")
        self.assertEqual(layout["margin_bottom"], "1.25")
        self.assertEqual(layout["margin_left"], "1.5")
        self.assertTrue(layout["width_px"])
        self.assertTrue(layout["height_px"])

    def test_render_docx_normalizes_inserted_value_color(self):
        with TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "template.docx"
            output_path = Path(temp_dir) / "output.docx"

            document = WordDocument()
            paragraph = document.add_paragraph("Client: ")
            run = paragraph.add_run("{{ full_name }}")
            run.font.color.rgb = RGBColor(112, 48, 160)
            document.save(template_path)

            DocxTemplateService.render_docx(
                template_path=str(template_path),
                output_path=str(output_path),
                values={"full_name": "Иван Иванов"},
            )

            rendered_document = WordDocument(output_path)
            rendered_runs = [
                run
                for paragraph in rendered_document.paragraphs
                for run in paragraph.runs
                if "Иван Иванов" in run.text
            ]

            self.assertTrue(rendered_runs)
            self.assertEqual(rendered_runs[0].font.color.rgb, RGBColor(0, 0, 0))

    def test_render_docx_uses_fast_replacements_for_plain_variables(self):
        with TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "template.docx"
            output_path = Path(temp_dir) / "output.docx"

            template = WordDocument()
            template.add_paragraph("Client: {{ full_name }}")
            template.add_paragraph("Program: {{ program_name }}")
            template.save(template_path)

            self.assertTrue(DocxTemplateService.can_render_with_simple_replacements(
                str(template_path),
            ))

            DocxTemplateService.render_docx(
                template_path=str(template_path),
                output_path=str(output_path),
                values={
                    "full_name": "A&B Client",
                    "program_name": "Law <Digital>",
                },
            )

            rendered_text = "\n".join(
                paragraph.text
                for paragraph in WordDocument(output_path).paragraphs
            )

            self.assertIn("A&B Client", rendered_text)
            self.assertIn("Law <Digital>", rendered_text)
            self.assertEqual(DocxTemplateService.extract_variables(str(output_path)), [])

    def test_add_page_numbering_adds_word_fields(self):
        with TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "numbered.docx"
            docx_path.write_bytes(self.create_docx_bytes("Contract body"))

            DocxTemplateService.add_page_numbering(str(docx_path))

            rendered_document = WordDocument(docx_path)
            footer_xml = rendered_document.sections[0].footer._element.xml
            self.assertIn("PAGE", footer_xml)
            self.assertIn("NUMPAGES", footer_xml)

    def test_prepare_editor_html_extracts_body_and_scopes_styles(self):
        html = """<!doctype html>
<html>
<head>
    <style>
        @page { size: A4; margin-left: 1in }
        p, table td { margin-bottom: 0; color: #111827 }
        body { margin: 0 }
    </style>
    <script>alert("x")</script>
</head>
<body lang="kk-KZ" dir="ltr">
    <p>Hello contract</p>
</body>
</html>"""

        prepared = DocxPreviewService.prepare_editor_html(html)

        self.assertIn('class="q-docx-html-fragment"', prepared)
        self.assertIn('lang="kk-KZ"', prepared)
        self.assertIn("Hello contract", prepared)
        self.assertIn(".q-docx-html-fragment p", prepared)
        self.assertIn(".q-docx-html-fragment table td", prepared)
        self.assertNotIn("<html", prepared.lower())
        self.assertNotIn("<body", prepared.lower())
        self.assertNotIn("@page", prepared)
        self.assertNotIn("<script", prepared.lower())

    def test_resolve_soffice_path_accepts_executable_path(self):
        with TemporaryDirectory() as temp_dir:
            soffice_path = Path(temp_dir) / "soffice.exe"
            soffice_path.touch()

            resolved = TemplateFileService.resolve_soffice_path(str(soffice_path))

        self.assertEqual(resolved, str(soffice_path))

    def test_resolve_soffice_path_accepts_program_directory(self):
        with TemporaryDirectory() as temp_dir:
            program_dir = Path(temp_dir) / "program"
            program_dir.mkdir()
            soffice_path = program_dir / "soffice.exe"
            soffice_path.touch()

            resolved = TemplateFileService.resolve_soffice_path(str(program_dir))

        self.assertEqual(resolved, str(soffice_path))

    @staticmethod
    def create_docx_bytes(text="Hello {{ full_name }}"):
        buffer = BytesIO()
        document = WordDocument()
        document.add_paragraph(text)
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def create_docx_bytes_with_margins():
        buffer = BytesIO()
        document = WordDocument()
        section = document.sections[0]
        section.top_margin = Cm(1)
        section.right_margin = Cm(1.5)
        section.bottom_margin = Cm(1.25)
        section.left_margin = Cm(1.5)
        document.add_paragraph("Document with custom margins")
        document.save(buffer)
        return buffer.getvalue()


class MoneyAmountServiceTests(SimpleTestCase):
    def test_build_value_context_returns_ru_and_kk_words(self):
        values = MoneyAmountService.build_value_context("tuition_amount", "450000")

        self.assertEqual(values["tuition_amount"], "450 000")
        self.assertEqual(
            values["tuition_amount_words_ru"],
            "четыреста пятьдесят тысяч тенге",
        )
        self.assertEqual(
            values["tuition_amount_words_kk"],
            "төрт жүз елу мың теңге",
        )
        self.assertEqual(
            values["tuition_amount_full_ru"],
            "450 000 (четыреста пятьдесят тысяч тенге)",
        )
        self.assertEqual(
            values["tuition_amount_full_kk"],
            "450 000 (төрт жүз елу мың теңге)",
        )

    def test_parse_amount_accepts_spaces_and_zero_fraction(self):
        self.assertEqual(MoneyAmountService.parse_amount("1 200 000,00"), 1200000)

    def test_parse_amount_rejects_fractional_or_non_numeric_values(self):
        self.assertIsNone(MoneyAmountService.parse_amount("1200.50"))
        self.assertIsNone(MoneyAmountService.parse_amount("twelve"))


class TemplateUploadFlowTests(TestCase):
    def test_upload_form_rejects_pdf(self):
        user, organization = self.create_user_and_organization()
        form = DocumentTemplateUploadForm(
            data={
                "organization": str(organization.pk),
                "title": "PDF template",
                "status": "active",
            },
            files={
                "template_file": SimpleUploadedFile(
                    "template.pdf",
                    b"%PDF-1.4",
                    content_type="application/pdf",
                )
            },
        )

        self.assertFalse(form.is_valid())
        self.assertIn("template_file", form.errors)

    def test_docx_upload_stores_docx_template_and_variables(self):
        with TemporaryDirectory() as temp_dir, override_settings(MEDIA_ROOT=temp_dir):
            user, organization = self.create_user_and_organization()
            self.client.force_login(user)

            response = self.client.post(
                reverse("documents:template_upload"),
                {
                    "title": "DOCX template",
                    "status": "active",
                    "template_file": SimpleUploadedFile(
                        "template.docx",
                        self.create_docx_bytes("Hello {{ full_name }}"),
                        content_type=(
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document"
                        ),
                    ),
                },
            )

            self.assertRedirects(response, reverse("documents:template_list"))

            template = DocumentTemplate.objects.get(title="DOCX template")
            self.assertTrue(template.template_file.name.endswith(".docx"))
            self.assertEqual(template.organization, organization)
            self.assertEqual(template.body_template, "")
            self.assertEqual(template.variables, ["full_name"])

    def test_doc_upload_is_normalized_to_docx(self):
        with TemporaryDirectory() as temp_dir, override_settings(MEDIA_ROOT=temp_dir):
            user, organization = self.create_user_and_organization()
            template = DocumentTemplate.objects.create(
                organization=organization,
                created_by=user,
                title="DOC template",
            )
            template.template_file.save(
                "legacy.doc",
                ContentFile(b"legacy binary content"),
                save=True,
            )

            def fake_convert_doc_to_docx(*, file_path, output_path):
                Path(output_path).write_bytes(
                    self.create_docx_bytes("Converted {{ contract_number }}")
                )

            with patch.object(
                TemplateFileService,
                "convert_doc_to_docx",
                side_effect=fake_convert_doc_to_docx,
            ):
                converted = TemplateFileService.normalize_template_file_to_docx(template)

            template.refresh_from_db()

            self.assertTrue(converted)
            self.assertTrue(template.template_file.name.endswith(".docx"))
            self.assertIn("contract_number", TemplateFileService.extract_variables(
                template.template_file.path
            ))

    def test_template_edit_uses_onlyoffice_for_docx_template(self):
        with TemporaryDirectory() as temp_dir, override_settings(
            MEDIA_ROOT=temp_dir,
            ONLYOFFICE_SERVER_URL="http://onlyoffice.test",
            ONLYOFFICE_DJANGO_URL="http://django.test",
            ONLYOFFICE_JWT_SECRET="test-secret",
        ):
            user, organization = self.create_user_and_organization()
            template = DocumentTemplate.objects.create(
                organization=organization,
                created_by=user,
                title="DOCX template",
            )
            template.template_file.save(
                "template.docx",
                ContentFile(self.create_docx_bytes("Hello {{ full_name }}")),
                save=True,
            )
            self.client.force_login(user)

            response = self.client.get(reverse("documents:template_edit", args=[template.pk]))

            self.assertEqual(response.status_code, 200)
            self.assertContains(response, "http://onlyoffice.test/web-apps/apps/api/documents/api.js")
            self.assertContains(response, "templateOnlyOfficeEditor")
            self.assertContains(response, "DocsAPI.DocEditor")
            self.assertContains(response, "http://django.test/documents/")

    def test_template_onlyoffice_callback_saves_updated_template_and_variables(self):
        with TemporaryDirectory() as temp_dir, override_settings(
            MEDIA_ROOT=temp_dir,
            ONLYOFFICE_JWT_SECRET="test-secret",
        ):
            user, organization = self.create_user_and_organization()
            template = DocumentTemplate.objects.create(
                organization=organization,
                created_by=user,
                title="DOCX template",
            )
            template.template_file.save(
                "template.docx",
                ContentFile(self.create_docx_bytes("Hello {{ old_name }}")),
                save=True,
            )
            token = OnlyOfficeService.encode_token({
                "template_id": template.pk,
                "action": "callback_template",
            })
            updated_bytes = self.create_docx_bytes("Hello {{ new_name }}")

            with patch("documents.views.requests.get") as get_mock:
                get_mock.return_value = SimpleNamespace(
                    content=updated_bytes,
                    raise_for_status=lambda: None,
                )
                response = self.client.post(
                    reverse("documents:template_onlyoffice_callback", args=[template.pk]),
                    data=json.dumps({"status": 2, "url": "http://onlyoffice.test/template.docx"}),
                    content_type="application/json",
                    QUERY_STRING=f"token={token}",
                )

            self.assertEqual(response.status_code, 200)
            self.assertEqual(response.json(), {"error": 0})

            template.refresh_from_db()
            self.assertEqual(template.body_template, "")
            self.assertEqual(template.variables, ["new_name"])
            with template.template_file.open("rb") as saved_file:
                self.assertEqual(saved_file.read(), updated_bytes)

    def test_template_onlyoffice_save_updates_schema_and_requests_force_save(self):
        with TemporaryDirectory() as temp_dir, override_settings(MEDIA_ROOT=temp_dir):
            user, organization = self.create_user_and_organization()
            template = DocumentTemplate.objects.create(
                organization=organization,
                created_by=user,
                title="DOCX template",
            )
            template.template_file.save(
                "template.docx",
                ContentFile(self.create_docx_bytes("Hello {{ full_name }}")),
                save=True,
            )
            self.client.force_login(user)
            schema = [{"title": "Contract", "fields": [{"label": "Number", "key": "contract_number"}]}]

            with patch.object(OnlyOfficeService, "force_save_key", return_value={"error": 0}) as force_save:
                response = self.client.post(
                    reverse("documents:template_onlyoffice_save", args=[template.pk]),
                    {
                        "field_schema": json.dumps(schema),
                        "onlyoffice_key": "opened-key",
                    },
                    HTTP_X_REQUESTED_WITH="XMLHttpRequest",
                )

            self.assertEqual(response.status_code, 200)
            self.assertEqual(response.json()["ok"], True)
            force_save.assert_called_once_with("opened-key")

            template.refresh_from_db()
            self.assertEqual(template.field_schema, schema)
            self.assertEqual(template.variables, ["full_name", "contract_number"])

    def test_new_template_party_gets_optional_system_email_field(self):
        user, organization = self.create_user_and_organization()
        template = DocumentTemplate.objects.create(
            organization=organization,
            created_by=user,
            title="Template with signer",
        )

        party = TemplateParty.objects.create(
            template=template,
            title="Customer",
            variable_prefix="customer",
            is_signer=True,
        )

        email_field = TemplatePartyField.objects.get(
            party=party,
            variable_name="email",
        )
        self.assertEqual(email_field.field_type, TemplatePartyField.FieldType.EMAIL)
        self.assertTrue(email_field.is_system)
        self.assertFalse(email_field.is_required)

    def test_template_editor_shows_document_system_fields_without_parties(self):
        user, organization = self.create_user_and_organization()
        template = DocumentTemplate.objects.create(
            organization=organization,
            created_by=user,
            title="Template without parties",
        )
        self.client.force_login(user)

        response = self.client.get(reverse("documents:template_edit", args=[template.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "contract_number")
        self.assertContains(response, "contract_date")
        self.assertContains(response, "contract_year")
        self.assertContains(response, "university_name_ru")
        self.assertContains(response, "university_bin")
        self.assertContains(response, "university_account")
        self.assertContains(response, "{{ contract_number }}")
        self.assertContains(response, "{{ contract_date }}")
        self.assertContains(response, "{{ university_name_ru }}")

    @staticmethod
    def create_docx_bytes(text):
        buffer = BytesIO()
        document = WordDocument()
        document.add_paragraph(text)
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def create_user_and_organization():
        user = get_user_model().objects.create_user(
            username="owner",
            password="password",
        )
        organization = Organization.objects.create(
            name="Test Organization",
            created_by=user,
        )
        OrganizationMember.objects.create(
            organization=organization,
            user=user,
            role=OrganizationMember.Role.OWNER,
        )
        return user, organization


class ObjectStorageServiceTests(TestCase):
    @override_settings(
        MINIO_BUCKET="test-bucket",
        MINIO_DEFAULT_RETENTION_DAYS=30,
        OBJECT_STORAGE_ENABLED=True,
    )
    def test_store_bytes_sets_minio_object_retention_for_uploaded_version(self):
        user, organization = self.create_user_and_organization()
        template = DocumentTemplate.objects.create(
            organization=organization,
            created_by=user,
            title="Storage template",
        )
        document = Document.objects.create(
            organization=organization,
            template=template,
            created_by=user,
            title="Storage document",
        )

        class FakeMinioClient:
            def __init__(self):
                self.retention_calls = []

            def put_object(self, *args, **kwargs):
                return SimpleNamespace(version_id="version-1", etag="etag-1")

            def set_object_retention(self, *args, **kwargs):
                self.retention_calls.append((args, kwargs))

        client = FakeMinioClient()

        with patch.object(ObjectStorageService, "_client", return_value=client):
            stored_object = ObjectStorageService.store_bytes(
                document=document,
                data=b"final document bytes",
                filename="final.pdf",
                content_type="application/pdf",
                object_type=StoredObject.ObjectType.FINAL_PDF,
                created_by=user,
            )

        self.assertEqual(stored_object.version_id, "version-1")
        self.assertEqual(stored_object.retention_mode, StoredObject.RetentionMode.COMPLIANCE)
        self.assertIsNotNone(stored_object.retention_until)
        self.assertEqual(len(client.retention_calls), 1)

        args, kwargs = client.retention_calls[0]
        self.assertEqual(args[0], "test-bucket")
        self.assertEqual(args[1], stored_object.object_key)
        self.assertEqual(args[2].mode, "COMPLIANCE")
        self.assertEqual(args[2].retain_until_date, stored_object.retention_until)
        self.assertEqual(kwargs["version_id"], "version-1")

    @staticmethod
    def create_user_and_organization():
        user = get_user_model().objects.create_user(
            username="storage-owner",
            password="password",
        )
        organization = Organization.objects.create(
            name="Storage Organization",
            created_by=user,
        )
        return user, organization


class DocumentLedgerServiceTests(TestCase):
    @override_settings(
        LEDGER_ENABLED=True,
        LEDGER_ACTOR="sign_app",
        LEDGER_EXTERNAL_ID_PREFIX="contract",
    )
    def test_submit_document_exports_pdf_and_stores_ledger_proof(self):
        user, organization = self.create_user_and_organization()
        document = self.create_signed_document(user, organization)
        pdf = ExportedPdf(
            filename="signed_contract.pdf",
            content=b"%PDF-1.4 test",
        )
        ledger_payload = {
            "id": "6642429a-7964-4a43-b32f-540dc4f55c25",
            "external_id": f"contract-{document.signed_at.year}-{document.pk:06d}",
            "document_hash": "d6b5ca52a83a74b695346d7936488a09b15274a075b468da490dad0e212b352a",
            "document_token": "document-token-123",
            "size_bytes": len(pdf.content),
            "created_at": "2026-07-10T12:00:25.318942Z",
            "ledger_proof": {
                "sequence": 1,
                "entry_hash": "e217d2d8cdaf7d0dc1e31fc3ee99",
                "previous_hash": "0" * 64,
                "server_signature_b64": "C0r4W",
                "server_key_id": "ed25519:2b0b31cae49ac85f4e849126",
                "created_at": "2026-07-10T12:00:25.318942Z",
            },
        }
        stored_pdf = self.create_stored_ledger_pdf(document, pdf.content)

        with patch(
            "documents.services.document_ledger_service.OnlyOfficePdfExportService.export_document_pdf",
            return_value=pdf,
        ) as export_document_pdf, patch(
            "documents.services.document_ledger_service.ObjectStorageService.store_bytes",
            return_value=stored_pdf,
        ) as store_bytes, patch(
            "documents.services.document_ledger_service.LedgerClient.submit_document",
            return_value=ledger_payload,
        ) as submit_document:
            record, cached = DocumentLedgerService.submit_document(
                document=document,
                requested_by=user,
            )

        self.assertFalse(cached)
        export_document_pdf.assert_called_once_with(document)
        store_bytes.assert_called_once()
        submit_document.assert_called_once()

        _, storage_kwargs = store_bytes.call_args
        self.assertEqual(storage_kwargs["document"], document)
        self.assertEqual(storage_kwargs["data"], pdf.content)
        self.assertEqual(storage_kwargs["filename"], pdf.filename)
        self.assertEqual(storage_kwargs["content_type"], "application/pdf")
        self.assertEqual(storage_kwargs["object_type"], StoredObject.ObjectType.LEDGER_PDF)

        _, kwargs = submit_document.call_args
        self.assertEqual(kwargs["filename"], "signed_contract.pdf")
        self.assertEqual(kwargs["content"], pdf.content)
        self.assertEqual(kwargs["actor"], "sign_app")
        self.assertEqual(kwargs["external_id"], ledger_payload["external_id"])
        metadata = json.loads(kwargs["metadata_json"])
        self.assertEqual(metadata["document_id"], document.pk)
        self.assertEqual(metadata["status"], Document.Status.SIGNED)

        record.refresh_from_db()
        self.assertEqual(record.status, DocumentLedgerRecord.Status.SUBMITTED)
        self.assertEqual(record.ledger_id, ledger_payload["id"])
        self.assertEqual(record.document_token, ledger_payload["document_token"])
        self.assertEqual(record.document_hash, ledger_payload["document_hash"])
        self.assertEqual(record.ledger_pdf_object, stored_pdf)
        self.assertEqual(record.sequence, 1)
        self.assertEqual(record.entry_hash, "e217d2d8cdaf7d0dc1e31fc3ee99")

    @override_settings(LEDGER_ENABLED=False)
    def test_submit_document_fails_before_pdf_export_when_ledger_disabled(self):
        user, organization = self.create_user_and_organization()
        document = self.create_signed_document(user, organization)

        with patch(
            "documents.services.document_ledger_service.OnlyOfficePdfExportService.export_document_pdf",
        ) as export_document_pdf:
            with self.assertRaises(DocumentLedgerError):
                DocumentLedgerService.submit_document(document=document, requested_by=user)

        export_document_pdf.assert_not_called()
        self.assertFalse(DocumentLedgerRecord.objects.exists())

    def test_document_list_links_to_ledger_proof_for_submitted_document(self):
        user, organization = self.create_user_and_organization()
        document = self.create_signed_document(user, organization)
        DocumentLedgerRecord.objects.create(
            document=document,
            status=DocumentLedgerRecord.Status.SUBMITTED,
            actor="sign_app",
            external_id=f"contract-{document.signed_at.year}-{document.pk:06d}",
            ledger_id="ledger-id",
            sequence=1,
        )
        self.client.force_login(user)

        response = self.client.get(reverse("documents:document_list"))

        self.assertContains(
            response,
            reverse("documents:document_ledger_proof", args=[document.pk]),
        )
        self.assertContains(response, "Integrity check")
        self.assertNotContains(response, "Send to ledger")

    def test_verify_record_checks_stored_pdf_hash_and_ledger_chain(self):
        user, organization = self.create_user_and_organization()
        document = self.create_signed_document(user, organization)
        pdf_content = b"%PDF-1.4 verified"
        pdf_hash = hashlib.sha256(pdf_content).hexdigest()
        stored_pdf = self.create_stored_ledger_pdf(document, pdf_content)
        record = DocumentLedgerRecord.objects.create(
            document=document,
            status=DocumentLedgerRecord.Status.SUBMITTED,
            actor="sign_app",
            external_id=f"contract-{document.signed_at.year}-{document.pk:06d}",
            ledger_id="ledger-id",
            document_hash=pdf_hash,
            ledger_pdf_object=stored_pdf,
            document_token="token",
        )

        with patch(
            "documents.services.document_ledger_service.ObjectStorageService.get_stored_object_bytes",
            return_value=pdf_content,
        ) as get_stored_object_bytes, patch(
            "documents.services.document_ledger_service.LedgerClient.verify",
            return_value={"ok": True, "deep": True},
        ) as verify_ledger:
            result = DocumentLedgerService.verify_record(record)

        get_stored_object_bytes.assert_called_once_with(stored_pdf)
        verify_ledger.assert_called_once_with(deep=True)
        self.assertTrue(result["hash_matches"])
        self.assertTrue(result["ledger_chain_ok"])

        record.refresh_from_db()
        self.assertEqual(
            record.last_verification_status,
            DocumentLedgerRecord.VerificationStatus.PASSED,
        )

    @staticmethod
    def create_signed_document(user, organization):
        return Document.objects.create(
            organization=organization,
            created_by=user,
            title="Signed ledger document",
            status=Document.Status.SIGNED,
            signed_at=timezone.now(),
            content_hash="a" * 64,
        )

    @staticmethod
    def create_stored_ledger_pdf(document, content):
        sha256 = hashlib.sha256(content).hexdigest()
        return StoredObject.objects.create(
            document=document,
            object_type=StoredObject.ObjectType.LEDGER_PDF,
            bucket="test-bucket",
            object_key=f"documents/{document.pk}/ledger_pdf/{sha256}-signed_contract.pdf",
            version_id="version-1",
            sha256=sha256,
            content_type="application/pdf",
            size_bytes=len(content),
        )

    @staticmethod
    def create_user_and_organization():
        user = get_user_model().objects.create_user(
            username="ledger-owner",
            password="password",
        )
        organization = Organization.objects.create(
            name="Ledger Test Organization",
            created_by=user,
        )
        OrganizationMember.objects.create(
            organization=organization,
            user=user,
            role=OrganizationMember.Role.OWNER,
        )
        return user, organization


class DocumentDraftEditingTests(TestCase):
    def test_document_gets_unique_contract_number_and_date(self):
        user, organization = self.create_user_and_organization()

        first_document = Document.objects.create(
            organization=organization,
            created_by=user,
            title="First document",
        )
        second_document = Document.objects.create(
            organization=organization,
            created_by=user,
            title="Second document",
        )

        self.assertIsNotNone(first_document.contract_date)
        self.assertEqual(first_document.contract_number, str(first_document.pk))
        self.assertNotEqual(
            first_document.contract_number,
            second_document.contract_number,
        )
        self.assertEqual(
            first_document.get_contract_system_values()["contract_date"],
            first_document.contract_date.strftime("%d.%m.%Y"),
        )
        self.assertRegex(
            first_document.get_contract_system_values()["contract_date_text_ru"],
            rf"^«{first_document.contract_date.day:02d}» .+ {first_document.contract_date.year} г\.$",
        )
        self.assertRegex(
            first_document.get_contract_system_values()["contract_date_text_kk"],
            rf"^«{first_document.contract_date.day:02d}» .+ {first_document.contract_date.year} ж\.$",
        )
        self.assertTrue(first_document.verification_token)
        self.assertNotEqual(
            first_document.verification_token,
            second_document.verification_token,
        )

    def test_document_skips_taken_contract_number(self):
        user, organization = self.create_user_and_organization()

        first_document = Document.objects.create(
            organization=organization,
            created_by=user,
            title="First document",
        )
        taken_number = str(first_document.pk + 1)
        first_document.contract_number = taken_number
        first_document.save(update_fields=["contract_number"])

        second_document = Document.objects.create(
            organization=organization,
            created_by=user,
            title="Second document",
        )

        self.assertNotEqual(second_document.contract_number, taken_number)
        self.assertFalse(
            Document.objects
            .exclude(pk=second_document.pk)
            .filter(contract_number=second_document.contract_number)
            .exists()
        )

    def test_document_fill_uses_contract_system_values_without_editable_fields(self):
        user, organization = self.create_user_and_organization()
        template = DocumentTemplate.objects.create(
            organization=organization,
            created_by=user,
            title="Contract identity template",
            body_template=(
                "No {{ contract_number }} from {{ contract_date }} "
                "alias {{ date }} year {{ contract_year }} "
                "{{ university_name_ru }} {{ university_bin }} for {{ customer_name }}"
            ),
            variables=[
                "contract_number",
                "contract_date",
                "date",
                "contract_year",
                "university_name_ru",
                "university_bin",
                "customer_name",
            ],
            field_schema=[
                {
                    "title": "Customer",
                    "fields": [
                        {
                            "label": "Customer name",
                            "key": "customer_name",
                            "placeholder": "Customer name",
                        }
                    ],
                }
            ],
        )
        document = self.create_document(user, organization, template)
        customer_field = DocumentFieldValue.objects.create(
            document=document,
            field_name="customer_name",
        )
        self.client.force_login(user)

        response = self.client.get(reverse("documents:document_fill", args=[document.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'data-variable="contract_number"')
        self.assertContains(response, f'value="{document.contract_number}"')
        self.assertNotIn(
            "contract_number",
            [field.field_name for field in response.context["fields"]],
        )
        self.assertNotIn(
            "university_name_ru",
            [field.field_name for field in response.context["fields"]],
        )
        self.assertContains(response, "q-fill-system-value")
        self.assertContains(response, 'key.startsWith("university_")')

        with patch.object(DocumentDocxRenderService, "render", return_value=document):
            response = self.client.post(
                reverse("documents:document_fill", args=[document.pk]),
                {
                    f"field_{customer_field.id}": "Test Customer",
                },
            )

        self.assertRedirects(
            response,
            f"{reverse('documents:document_list')}?signers={document.pk}",
        )
        document.refresh_from_db()
        self.assertIn(document.contract_number, document.rendered_html)
        self.assertIn(document.get_contract_date_display(), document.rendered_html)
        self.assertIn(str(document.contract_date.year), document.rendered_html)
        self.assertIn(
            Document.UNIVERSITY_SYSTEM_FIELD_DEFAULTS["university_name_ru"],
            document.rendered_html,
        )
        self.assertIn(
            Document.UNIVERSITY_SYSTEM_FIELD_DEFAULTS["university_bin"],
            document.rendered_html,
        )
        self.assertIn("Test Customer", document.rendered_html)

    def test_document_fill_expands_money_amount_variables(self):
        user, organization = self.create_user_and_organization()
        template = DocumentTemplate.objects.create(
            organization=organization,
            created_by=user,
            title="Tuition template",
            body_template=(
                "{{ tuition_amount }} "
                "{{ tuition_amount_full_ru }} "
                "{{ tuition_amount_full_kk }}"
            ),
            variables=[
                "tuition_amount",
                "tuition_amount_full_ru",
                "tuition_amount_full_kk",
            ],
            field_schema=[
                {
                    "title": "Tuition",
                    "fields": [
                        {
                            "label": "Tuition amount",
                            "key": "tuition_amount",
                            "type": "money",
                            "placeholder": "Tuition amount",
                        }
                    ],
                }
            ],
        )
        document = self.create_document(user, organization, template)
        amount_field = DocumentFieldValue.objects.create(
            document=document,
            field_name="tuition_amount",
        )
        DocumentFieldValue.objects.create(
            document=document,
            field_name="tuition_amount_full_ru",
        )
        self.client.force_login(user)

        response = self.client.get(reverse("documents:document_fill", args=[document.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'data-field-type="money"')
        self.assertIn(
            "tuition_amount",
            [field.field_name for field in response.context["fields"]],
        )
        self.assertNotIn(
            "tuition_amount_full_ru",
            [field.field_name for field in response.context["fields"]],
        )

        with patch.object(DocumentDocxRenderService, "render", return_value=document):
            response = self.client.post(
                reverse("documents:document_fill", args=[document.pk]),
                {
                    f"field_{amount_field.id}": "450000",
                },
            )

        self.assertRedirects(
            response,
            f"{reverse('documents:document_list')}?signers={document.pk}",
        )
        document.refresh_from_db()
        self.assertIn("450 000", document.rendered_html)
        self.assertIn("четыреста пятьдесят тысяч тенге", document.rendered_html)
        self.assertIn("төрт жүз елу мың теңге", document.rendered_html)

    def test_docx_render_includes_contract_system_values(self):
        with TemporaryDirectory() as temp_dir, override_settings(MEDIA_ROOT=temp_dir):
            user, organization = self.create_user_and_organization()
            template = DocumentTemplate.objects.create(
                organization=organization,
                created_by=user,
                title="DOCX contract identity template",
            )
            template.template_file.save(
                "template.docx",
                ContentFile(self.create_docx_bytes("No {{ contract_number }}")),
                save=True,
            )
            document = self.create_document(user, organization, template)
            DocumentFieldValue.objects.create(
                document=document,
                field_name="customer_name",
                field_value="Test Customer",
            )
            captured_values = {}

            def fake_render_docx(*, template_path, output_path, values):
                captured_values.update(values)
                Path(output_path).write_bytes(b"rendered docx")

            with patch(
                "documents.services.document_docx_render_service.DocxTemplateService.render_docx",
                side_effect=fake_render_docx,
            ), patch(
                "documents.services.document_docx_render_service.DocumentVerificationAppendixService.finalize_docx"
            ):
                DocumentDocxRenderService.render(document)

            self.assertEqual(
                captured_values["contract_number"],
                document.contract_number,
            )
            self.assertEqual(
                captured_values["contract_date"],
                document.get_contract_date_display(),
            )
            self.assertEqual(
                captured_values["contract_date_text_ru"],
                document.get_contract_date_text_ru(),
            )
            self.assertEqual(
                captured_values["contract_date_text_kk"],
                document.get_contract_date_text_kk(),
            )
            self.assertEqual(
                captured_values["date"],
                document.get_contract_date_display(),
            )
            self.assertEqual(
                captured_values["contract_year"],
                str(document.contract_date.year),
            )
            self.assertEqual(
                captured_values["university_name_ru"],
                Document.UNIVERSITY_SYSTEM_FIELD_DEFAULTS["university_name_ru"],
            )
            self.assertEqual(
                captured_values["university_bin"],
                Document.UNIVERSITY_SYSTEM_FIELD_DEFAULTS["university_bin"],
            )

    def test_docx_render_includes_money_amount_variables(self):
        with TemporaryDirectory() as temp_dir, override_settings(MEDIA_ROOT=temp_dir):
            user, organization = self.create_user_and_organization()
            template = DocumentTemplate.objects.create(
                organization=organization,
                created_by=user,
                title="Money DOCX template",
                field_schema=[
                    {
                        "title": "Tuition",
                        "fields": [
                            {
                                "label": "Tuition amount",
                                "key": "tuition_amount",
                                "type": "money",
                                "placeholder": "Tuition amount",
                            }
                        ],
                    }
                ],
            )
            template.template_file.save(
                "template.docx",
                ContentFile(self.create_docx_bytes("{{ tuition_amount_full_ru }}")),
                save=True,
            )
            document = self.create_document(user, organization, template)
            DocumentFieldValue.objects.create(
                document=document,
                field_name="tuition_amount",
                field_value="450000",
            )
            captured_values = {}

            def fake_render_docx(*, template_path, output_path, values):
                captured_values.update(values)
                Path(output_path).write_bytes(b"rendered docx")

            with patch(
                "documents.services.document_docx_render_service.DocxTemplateService.render_docx",
                side_effect=fake_render_docx,
            ), patch(
                "documents.services.document_docx_render_service.DocumentVerificationAppendixService.finalize_docx"
            ):
                DocumentDocxRenderService.render(document)

            self.assertEqual(captured_values["tuition_amount"], "450 000")
            self.assertEqual(
                captured_values["tuition_amount_words_ru"],
                "четыреста пятьдесят тысяч тенге",
            )
            self.assertEqual(
                captured_values["tuition_amount_words_kk"],
                "төрт жүз елу мың теңге",
            )
            self.assertEqual(
                captured_values["tuition_amount_full_ru"],
                "450 000 (четыреста пятьдесят тысяч тенге)",
            )
            self.assertEqual(
                captured_values["tuition_amount_full_kk"],
                "450 000 (төрт жүз елу мың теңге)",
            )

    def test_verification_appendix_is_added_once_to_docx(self):
        with TemporaryDirectory() as temp_dir, override_settings(
            MEDIA_ROOT=temp_dir,
            PUBLIC_SITE_URL="https://qolqoyu.example.test",
        ):
            user, organization = self.create_user_and_organization()
            template = self.create_template(user, organization)
            document = self.create_document(
                user,
                organization,
                template,
                title="Appendix document",
            )
            docx_path = Path(temp_dir) / "appendix.docx"
            docx_path.write_bytes(self.create_docx_bytes("Document body"))

            DocumentVerificationAppendixService.finalize_docx(
                document=document,
                file_path=str(docx_path),
            )
            DocumentVerificationAppendixService.finalize_docx(
                document=document,
                file_path=str(docx_path),
            )

            rendered_document = WordDocument(docx_path)
            text_parts = [paragraph.text for paragraph in rendered_document.paragraphs]
            for table in rendered_document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            full_text = "\n".join(text_parts)

            self.assertEqual(full_text.count("Проверка документа"), 1)
            self.assertIn(document.contract_number, full_text)
            self.assertIn(
                f"https://qolqoyu.example.test/documents/verify/{document.verification_token}/",
                full_text,
            )
            self.assertIn("QR-код", full_text)

    def test_create_document_from_uploaded_docx_without_template(self):
        with TemporaryDirectory() as temp_dir, override_settings(MEDIA_ROOT=temp_dir):
            user, organization = self.create_user_and_organization()
            self.client.force_login(user)

            response = self.client.post(
                reverse("documents:document_create"),
                {
                    "title": "Uploaded document",
                    "document_file": SimpleUploadedFile(
                        "ready.docx",
                        self.create_docx_bytes("Ready document"),
                        content_type=(
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document"
                        ),
                    ),
                },
            )

            document = Document.objects.get(title="Uploaded document")

            self.assertRedirects(
                response,
                reverse("documents:document_onlyoffice_editor", args=[document.pk]),
            )
            self.assertEqual(document.organization, organization)
            self.assertIsNone(document.template)
            self.assertTrue(document.rendered_docx_file.name.endswith(".docx"))
            self.assertTrue(document.content_hash)

    def test_create_document_from_uploaded_doc_is_converted_to_docx(self):
        with TemporaryDirectory() as temp_dir, override_settings(MEDIA_ROOT=temp_dir):
            user, organization = self.create_user_and_organization()
            self.client.force_login(user)

            def fake_convert_doc_to_docx(*, file_path, output_path):
                Path(output_path).write_bytes(self.create_docx_bytes("Converted document"))

            with patch.object(
                TemplateFileService,
                "convert_doc_to_docx",
                side_effect=fake_convert_doc_to_docx,
            ):
                response = self.client.post(
                    reverse("documents:document_create"),
                    {
                        "title": "Uploaded legacy document",
                        "document_file": SimpleUploadedFile(
                            "legacy.doc",
                            b"legacy binary content",
                            content_type="application/msword",
                        ),
                    },
                )

            document = Document.objects.get(title="Uploaded legacy document")

            self.assertRedirects(
                response,
                reverse("documents:document_onlyoffice_editor", args=[document.pk]),
            )
            self.assertEqual(document.organization, organization)
            self.assertIsNone(document.template)
            self.assertTrue(document.rendered_docx_file.name.endswith(".docx"))
            self.assertTrue(document.content_hash)

    def test_uploaded_document_editor_redirects_to_onlyoffice(self):
        with TemporaryDirectory() as temp_dir, override_settings(MEDIA_ROOT=temp_dir):
            user, organization = self.create_user_and_organization()
            document = Document.objects.create(
                organization=organization,
                created_by=user,
                title="Uploaded document",
                status=Document.Status.DRAFT,
            )
            document.rendered_docx_file.save(
                "ready.docx",
                ContentFile(self.create_docx_bytes("Ready document")),
                save=True,
            )
            self.client.force_login(user)

            response = self.client.get(reverse("documents:document_editor", args=[document.pk]))

            self.assertRedirects(response, reverse("documents:document_onlyoffice_editor", args=[document.pk]))

    def test_uploaded_document_onlyoffice_editor_builds_config(self):
        with TemporaryDirectory() as temp_dir, override_settings(
            MEDIA_ROOT=temp_dir,
            ONLYOFFICE_SERVER_URL="http://onlyoffice.test",
            ONLYOFFICE_DJANGO_URL="http://django.test",
            ONLYOFFICE_JWT_SECRET="test-secret",
        ):
            user, organization = self.create_user_and_organization()
            document = Document.objects.create(
                organization=organization,
                created_by=user,
                title="Uploaded document",
                status=Document.Status.DRAFT,
            )
            document.rendered_docx_file.save(
                "ready.docx",
                ContentFile(self.create_docx_bytes("Ready document")),
                save=True,
            )
            self.client.force_login(user)

            response = self.client.get(reverse("documents:document_onlyoffice_editor", args=[document.pk]))

            self.assertEqual(response.status_code, 200)
            self.assertContains(response, "http://onlyoffice.test/web-apps/apps/api/documents/api.js")
            self.assertContains(response, "DocsAPI.DocEditor")
            self.assertContains(response, "http://django.test/documents/")
            self.assertContains(response, '"documentType": "word"')

    @override_settings(ONLYOFFICE_COMMAND_SERVICE_URL="http://127.0.0.1:8082/command")
    def test_onlyoffice_convert_service_url_uses_modern_converter_endpoint(self):
        self.assertEqual(
            OnlyOfficeService.convert_service_url("convert-key"),
            "http://127.0.0.1:8082/converter?shardkey=convert-key",
        )

    def test_document_list_shows_edit_link_for_uploaded_draft_document(self):
        with TemporaryDirectory() as temp_dir, override_settings(MEDIA_ROOT=temp_dir):
            user, organization = self.create_user_and_organization()
            document = Document.objects.create(
                organization=organization,
                created_by=user,
                title="Uploaded document",
                status=Document.Status.DRAFT,
            )
            document.rendered_docx_file.save(
                "ready.docx",
                ContentFile(self.create_docx_bytes("Ready document")),
                save=True,
            )
            self.client.force_login(user)

            response = self.client.get(reverse("documents:document_list"))

            self.assertEqual(response.status_code, 200)
            self.assertContains(response, reverse("documents:document_onlyoffice_editor", args=[document.pk]))

    def test_onlyoffice_file_endpoint_returns_docx_with_valid_token(self):
        with TemporaryDirectory() as temp_dir, override_settings(
            MEDIA_ROOT=temp_dir,
            ONLYOFFICE_JWT_SECRET="test-secret",
        ):
            user, organization = self.create_user_and_organization()
            document = Document.objects.create(
                organization=organization,
                created_by=user,
                title="Uploaded document",
                status=Document.Status.DRAFT,
            )
            document.rendered_docx_file.save(
                "ready.docx",
                ContentFile(self.create_docx_bytes("Ready document")),
                save=True,
            )
            token = OnlyOfficeService.encode_token({
                "document_id": document.pk,
                "action": "download",
            })

            response = self.client.get(
                reverse("documents:document_onlyoffice_file", args=[document.pk]),
                {"token": token},
            )

            self.assertEqual(response.status_code, 200)
            self.assertEqual(
                response["Content-Type"],
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            self.assertTrue(response.content)

    def test_onlyoffice_callback_saves_updated_docx(self):
        with TemporaryDirectory() as temp_dir, override_settings(
            MEDIA_ROOT=temp_dir,
            ONLYOFFICE_JWT_SECRET="test-secret",
        ):
            user, organization = self.create_user_and_organization()
            document = Document.objects.create(
                organization=organization,
                created_by=user,
                title="Uploaded document",
                status=Document.Status.DRAFT,
            )
            document.rendered_docx_file.save(
                "ready.docx",
                ContentFile(self.create_docx_bytes("Ready document")),
                save=True,
            )
            original_docx_name = document.rendered_docx_file.name
            token = OnlyOfficeService.encode_token({
                "document_id": document.pk,
                "action": "callback",
            })
            updated_bytes = self.create_docx_bytes("Updated document")

            with patch("documents.views.requests.get") as get_mock:
                get_mock.return_value = SimpleNamespace(
                    content=updated_bytes,
                    raise_for_status=lambda: None,
                )
                response = self.client.post(
                    reverse("documents:document_onlyoffice_callback", args=[document.pk]),
                    data=json.dumps({"status": 2, "url": "http://onlyoffice.test/saved.docx"}),
                    content_type="application/json",
                    QUERY_STRING=f"token={token}",
                )

            self.assertEqual(response.status_code, 200)
            self.assertEqual(response.json(), {"error": 0})

            document.refresh_from_db()
            self.assertEqual(document.rendered_html, "")
            self.assertNotEqual(document.rendered_docx_file.name, original_docx_name)
            self.assertTrue(document.rendered_docx_file.name.endswith(".docx"))
            with document.rendered_docx_file.open("rb") as saved_file:
                saved_bytes = saved_file.read()

            self.assertNotEqual(saved_bytes, updated_bytes)
            saved_document = WordDocument(BytesIO(saved_bytes))
            saved_text = "\n".join(
                [paragraph.text for paragraph in saved_document.paragraphs]
                + [
                    cell.text
                    for table in saved_document.tables
                    for row in table.rows
                    for cell in row.cells
                ]
            )
            self.assertIn("Updated document", saved_text)
            self.assertIn("Проверка документа", saved_text)

    def test_onlyoffice_callback_ignores_non_save_status(self):
        with TemporaryDirectory() as temp_dir, override_settings(
            MEDIA_ROOT=temp_dir,
            ONLYOFFICE_JWT_SECRET="test-secret",
        ):
            user, organization = self.create_user_and_organization()
            document = Document.objects.create(
                organization=organization,
                created_by=user,
                title="Uploaded document",
                status=Document.Status.DRAFT,
            )
            document.rendered_docx_file.save(
                "ready.docx",
                ContentFile(self.create_docx_bytes("Ready document")),
                save=True,
            )
            original_docx_name = document.rendered_docx_file.name
            token = OnlyOfficeService.encode_token({
                "document_id": document.pk,
                "action": "callback",
            })

            response = self.client.post(
                reverse("documents:document_onlyoffice_callback", args=[document.pk]),
                data=json.dumps({"status": 1}),
                content_type="application/json",
                QUERY_STRING=f"token={token}",
            )

            document.refresh_from_db()
            self.assertEqual(response.json(), {"error": 0})
            self.assertEqual(document.rendered_docx_file.name, original_docx_name)

    def test_uploaded_document_fill_redirects_to_signers_modal(self):
        with TemporaryDirectory() as temp_dir, override_settings(MEDIA_ROOT=temp_dir):
            user, organization = self.create_user_and_organization()
            document = Document.objects.create(
                organization=organization,
                created_by=user,
                title="Uploaded document",
                status=Document.Status.DRAFT,
            )
            document.rendered_docx_file.save(
                "ready.docx",
                ContentFile(self.create_docx_bytes("Ready document")),
                save=True,
            )
            self.client.force_login(user)

            response = self.client.get(reverse("documents:document_fill", args=[document.pk]))

            self.assertRedirects(
                response,
                reverse("signing:document_signers", args=[document.pk]),
            )

    def test_uploaded_document_signers_page_adds_manual_signer(self):
        user, organization = self.create_user_and_organization()
        document = Document.objects.create(
            organization=organization,
            created_by=user,
            title="Uploaded document",
            status=Document.Status.DRAFT,
        )
        self.client.force_login(user)

        response = self.client.post(
            reverse("signing:document_signers", args=[document.pk]),
            {
                "full_name": "Manual Signer",
                "iin": "123456789012",
                "phone": "+77071234567",
                "signing_order": "1",
                "signing_method": "sms",
            },
        )

        self.assertRedirects(response, reverse("signing:document_signers", args=[document.pk]))
        signer = document.signers.get()
        self.assertEqual(signer.full_name, "Manual Signer")
        self.assertIsNone(signer.template_party)

    def test_document_list_shows_edit_link_until_first_signature(self):
        user, organization = self.create_user_and_organization()
        template = self.create_template(user, organization)
        draft_document = self.create_document(
            user,
            organization,
            template,
            title="Draft document",
            status=Document.Status.DRAFT,
        )
        locked_document = self.create_document(
            user,
            organization,
            template,
            title="Waiting document",
            status=Document.Status.WAITING_FOR_SIGNERS,
        )
        signed_document = self.create_document(
            user,
            organization,
            template,
            title="Partially signed document",
            status=Document.Status.WAITING_FOR_SIGNERS,
        )
        Signer.objects.create(
            document=signed_document,
            full_name="Signed Person",
            iin="123456789012",
            phone="77071234567",
            status=Signer.Status.SIGNED,
        )
        self.client.force_login(user)

        response = self.client.get(reverse("documents:document_list"))

        self.assertContains(
            response,
            reverse("documents:document_fill", args=[draft_document.pk]),
        )
        self.assertContains(
            response,
            reverse("documents:document_fill", args=[locked_document.pk]),
        )
        self.assertNotContains(
            response,
            reverse("documents:document_fill", args=[signed_document.pk]),
        )

    def test_document_list_shows_contract_number(self):
        user, organization = self.create_user_and_organization()
        template = self.create_template(user, organization)
        document = self.create_document(
            user,
            organization,
            template,
            title="Numbered document",
        )
        self.client.force_login(user)

        response = self.client.get(reverse("documents:document_list"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Contract number")
        self.assertContains(response, document.contract_number)

    def test_document_list_links_to_public_verification_page(self):
        user, organization = self.create_user_and_organization()
        template = self.create_template(user, organization)
        document = self.create_document(
            user,
            organization,
            template,
            title="Verifiable document",
        )
        self.client.force_login(user)

        response = self.client.get(reverse("documents:document_list"))

        self.assertContains(response, "Verification page")
        self.assertContains(
            response,
            reverse("documents:document_verification", args=[document.verification_token]),
        )

    def test_document_verification_page_is_public(self):
        user, organization = self.create_user_and_organization()
        template = self.create_template(user, organization)
        document = self.create_document(
            user,
            organization,
            template,
            title="Public verification document",
        )
        document.rendered_html = "Public verification content"
        document.update_content_hash(save=True)

        response = self.client.get(
            reverse("documents:document_verification", args=[document.verification_token])
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Document integrity")
        self.assertContains(response, document.contract_number)
        self.assertContains(response, document.content_hash)

    def test_document_fill_allows_waiting_document_without_signatures(self):
        user, organization = self.create_user_and_organization()
        template = self.create_template(user, organization)
        document = self.create_document(
            user,
            organization,
            template,
            status=Document.Status.WAITING_FOR_SIGNERS,
        )
        self.client.force_login(user)

        response = self.client.get(reverse("documents:document_fill", args=[document.pk]))

        self.assertEqual(response.status_code, 200)

    def test_document_fill_redirects_after_first_signature(self):
        user, organization = self.create_user_and_organization()
        template = self.create_template(user, organization)
        document = self.create_document(
            user,
            organization,
            template,
            status=Document.Status.WAITING_FOR_SIGNERS,
        )
        Signer.objects.create(
            document=document,
            full_name="Signed Person",
            iin="123456789012",
            phone="77071234567",
            status=Signer.Status.SIGNED,
        )
        self.client.force_login(user)

        response = self.client.get(reverse("documents:document_fill", args=[document.pk]))

        self.assertRedirects(response, reverse("documents:document_list"))

    def test_document_fill_handles_invalid_template_party_signer_data(self):
        user, organization = self.create_user_and_organization()
        template = self.create_template(user, organization)
        party = TemplateParty.objects.create(
            template=template,
            title="Customer",
            variable_prefix="customer",
            is_signer=True,
        )
        document = self.create_document(user, organization, template)
        for field_name in [
            "customer_full_name",
            "customer_iin_bin",
            "customer_phone",
        ]:
            DocumentFieldValue.objects.create(
                document=document,
                field_name=field_name,
            )
        self.client.force_login(user)
        field_inputs = {
            field.field_name: field.id
            for field in document.field_values.all()
        }

        response = self.client.post(
            reverse("documents:document_fill", args=[document.pk]),
            {
                f"field_{field_inputs['customer_full_name']}": "Test Signer",
                f"field_{field_inputs['customer_iin_bin']}": "123",
                f"field_{field_inputs['customer_phone']}": "+77071234567",
            },
        )

        self.assertRedirects(response, reverse("documents:document_fill", args=[document.pk]))
        messages = list(response.wsgi_request._messages)
        self.assertTrue(any("ИИН должен содержать ровно 12 цифр" in str(item) for item in messages))

    def test_document_fill_creates_template_signer_with_empty_email(self):
        user, organization = self.create_user_and_organization()
        template = self.create_template(user, organization)
        party = TemplateParty.objects.create(
            template=template,
            title="Customer",
            variable_prefix="customer",
            is_signer=True,
        )
        document = self.create_document(user, organization, template)
        self.create_signer_field_values(document, party)
        self.client.force_login(user)
        field_inputs = {
            field.field_name: field.id
            for field in document.field_values.all()
        }

        response = self.client.post(
            reverse("documents:document_fill", args=[document.pk]),
            {
                f"field_{field_inputs['customer_full_name']}": "Test Signer",
                f"field_{field_inputs['customer_iin_bin']}": "123456789012",
                f"field_{field_inputs['customer_phone']}": "+77071234567",
                f"field_{field_inputs['customer_email']}": "",
            },
        )

        self.assertRedirects(response, f"{reverse('documents:document_list')}?signers={document.pk}")
        signer = document.signers.get()
        self.assertEqual(signer.email, "")

    def test_document_fill_creates_template_signer_with_valid_email(self):
        user, organization = self.create_user_and_organization()
        template = self.create_template(user, organization)
        party = TemplateParty.objects.create(
            template=template,
            title="Customer",
            variable_prefix="customer",
            is_signer=True,
        )
        document = self.create_document(user, organization, template)
        self.create_signer_field_values(document, party)
        self.client.force_login(user)
        field_inputs = {
            field.field_name: field.id
            for field in document.field_values.all()
        }

        with patch("signing.services.signer_service.SignerService.email_domain_exists", return_value=True):
            response = self.client.post(
                reverse("documents:document_fill", args=[document.pk]),
                {
                    f"field_{field_inputs['customer_full_name']}": "Test Signer",
                    f"field_{field_inputs['customer_iin_bin']}": "123456789012",
                    f"field_{field_inputs['customer_phone']}": "+77071234567",
                    f"field_{field_inputs['customer_email']}": "signer@example.com",
                },
            )

        self.assertRedirects(response, f"{reverse('documents:document_list')}?signers={document.pk}")
        signer = document.signers.get()
        self.assertEqual(signer.email, "signer@example.com")

    def test_document_fill_rejects_invalid_template_signer_email(self):
        user, organization = self.create_user_and_organization()
        template = self.create_template(user, organization)
        party = TemplateParty.objects.create(
            template=template,
            title="Customer",
            variable_prefix="customer",
            is_signer=True,
        )
        document = self.create_document(user, organization, template)
        self.create_signer_field_values(document, party)
        self.client.force_login(user)
        field_inputs = {
            field.field_name: field.id
            for field in document.field_values.all()
        }

        response = self.client.post(
            reverse("documents:document_fill", args=[document.pk]),
            {
                f"field_{field_inputs['customer_full_name']}": "Test Signer",
                f"field_{field_inputs['customer_iin_bin']}": "123456789012",
                f"field_{field_inputs['customer_phone']}": "+77071234567",
                f"field_{field_inputs['customer_email']}": "not-an-email",
            },
        )

        self.assertRedirects(response, reverse("documents:document_fill", args=[document.pk]))
        self.assertFalse(document.signers.exists())
        messages = list(response.wsgi_request._messages)
        self.assertTrue(any("Email must be a valid email address." in str(item) for item in messages))

    def test_document_fill_rejects_unverified_template_signer_email_domain(self):
        user, organization = self.create_user_and_organization()
        template = self.create_template(user, organization)
        party = TemplateParty.objects.create(
            template=template,
            title="Customer",
            variable_prefix="customer",
            is_signer=True,
        )
        document = self.create_document(user, organization, template)
        self.create_signer_field_values(document, party)
        self.client.force_login(user)
        field_inputs = {
            field.field_name: field.id
            for field in document.field_values.all()
        }

        with patch("signing.services.signer_service.SignerService.email_domain_exists", return_value=False):
            response = self.client.post(
                reverse("documents:document_fill", args=[document.pk]),
                {
                    f"field_{field_inputs['customer_full_name']}": "Test Signer",
                    f"field_{field_inputs['customer_iin_bin']}": "123456789012",
                    f"field_{field_inputs['customer_phone']}": "+77071234567",
                    f"field_{field_inputs['customer_email']}": "signer@missing.example",
                },
            )

        self.assertRedirects(response, reverse("documents:document_fill", args=[document.pk]))
        self.assertFalse(document.signers.exists())
        messages = list(response.wsgi_request._messages)
        self.assertTrue(any("Email domain could not be verified." in str(item) for item in messages))

    @staticmethod
    def create_template(user, organization):
        return DocumentTemplate.objects.create(
            organization=organization,
            created_by=user,
            title="Editable template",
            body_template=(
                "Hello {{ full_name }} {{ customer_full_name }} "
                "{{ customer_iin_bin }} {{ customer_phone }}"
            ),
            variables=[
                "full_name",
                "customer_full_name",
                "customer_iin_bin",
                "customer_phone",
            ],
        )

    @staticmethod
    def create_document(user, organization, template, title="Document", status=Document.Status.DRAFT):
        return Document.objects.create(
            organization=organization,
            template=template,
            created_by=user,
            title=title,
            status=status,
        )

    @staticmethod
    def create_signer_field_values(document, party):
        for field in party.fields.all():
            DocumentFieldValue.objects.create(
                document=document,
                field_name=f"{party.variable_prefix}_{field.variable_name}",
            )

    @staticmethod
    def create_docx_bytes(text):
        buffer = BytesIO()
        document = WordDocument()
        document.add_paragraph(text)
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def create_docx_bytes_with_margins(text):
        buffer = BytesIO()
        document = WordDocument()
        section = document.sections[0]
        section.top_margin = Cm(1)
        section.right_margin = Cm(1.5)
        section.bottom_margin = Cm(1.25)
        section.left_margin = Cm(1.5)
        document.add_paragraph(text)
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def create_user_and_organization():
        user = get_user_model().objects.create_user(
            username="document-owner",
            password="password",
        )
        organization = Organization.objects.create(
            name="Document Test Organization",
            created_by=user,
        )
        OrganizationMember.objects.create(
            organization=organization,
            user=user,
            role=OrganizationMember.Role.OWNER,
        )
        return user, organization


class OrganizationAccessTests(TestCase):
    def test_organization_admin_sees_templates_and_documents_created_by_another_user(self):
        owner = self.create_user("owner")
        admin = self.create_user("admin")
        creator = self.create_user("creator")
        organization = self.create_organization(owner)
        OrganizationMember.objects.create(
            organization=organization,
            user=admin,
            role=OrganizationMember.Role.ADMIN,
        )
        template = self.create_template(creator, organization, title="Shared template")
        document = self.create_document(creator, organization, template, title="Shared document")
        self.client.force_login(admin)

        template_list_response = self.client.get(reverse("documents:template_list"))
        document_list_response = self.client.get(reverse("documents:document_list"))

        self.assertContains(template_list_response, "Shared template")
        self.assertContains(document_list_response, "Shared document")
        self.assertEqual(
            self.client.get(reverse("documents:template_edit", args=[template.pk])).status_code,
            200,
        )
        self.assertEqual(
            self.client.get(reverse("documents:document_fill", args=[document.pk])).status_code,
            200,
        )

    def test_organization_member_sees_only_own_department_templates_and_documents(self):
        owner = self.create_user("owner")
        member = self.create_user("member")
        organization = self.create_organization(owner)
        own_department = Department.objects.create(
            organization=organization,
            name="Own department",
        )
        other_department = Department.objects.create(
            organization=organization,
            name="Other department",
        )
        OrganizationMember.objects.create(
            organization=organization,
            user=member,
            role=OrganizationMember.Role.MEMBER,
            department=own_department,
        )
        own_template = self.create_template(
            owner,
            organization,
            department=own_department,
            title="Own department template",
        )
        other_template = self.create_template(
            owner,
            organization,
            department=other_department,
            title="Other department template",
        )
        own_document = self.create_document(
            owner,
            organization,
            own_template,
            title="Own department document",
        )
        other_document = self.create_document(
            owner,
            organization,
            other_template,
            title="Other department document",
        )
        self.client.force_login(member)

        template_list_response = self.client.get(reverse("documents:template_list"))
        document_list_response = self.client.get(reverse("documents:document_list"))

        self.assertContains(template_list_response, "Own department template")
        self.assertNotContains(template_list_response, "Other department template")
        self.assertContains(document_list_response, "Own department document")
        self.assertNotContains(document_list_response, "Other department document")
        self.assertEqual(
            self.client.get(reverse("documents:template_edit", args=[own_template.pk])).status_code,
            200,
        )
        self.assertEqual(
            self.client.get(reverse("documents:template_edit", args=[other_template.pk])).status_code,
            404,
        )
        self.assertEqual(
            self.client.get(reverse("documents:document_fill", args=[own_document.pk])).status_code,
            200,
        )
        self.assertEqual(
            self.client.get(reverse("documents:document_fill", args=[other_document.pk])).status_code,
            404,
        )

    def test_foreign_admin_cannot_access_other_organization_documents(self):
        owner = self.create_user("owner")
        foreign_owner = self.create_user("foreign-owner")
        organization = self.create_organization(owner, name="Owner Org")
        foreign_organization = self.create_organization(foreign_owner, name="Foreign Org")
        template = self.create_template(owner, organization)
        document = self.create_document(owner, organization, template)
        self.client.force_login(foreign_owner)

        self.assertNotContains(
            self.client.get(reverse("documents:document_list")),
            document.title,
        )
        self.assertEqual(
            self.client.get(reverse("documents:document_fill", args=[document.pk])).status_code,
            404,
        )
        self.assertNotEqual(organization.pk, foreign_organization.pk)

    @staticmethod
    def create_user(username):
        return get_user_model().objects.create_user(
            username=username,
            password="password",
        )

    @staticmethod
    def create_organization(user, name="Access Organization"):
        organization = Organization.objects.create(
            name=name,
            created_by=user,
        )
        OrganizationMember.objects.create(
            organization=organization,
            user=user,
            role=OrganizationMember.Role.OWNER,
        )
        return organization

    @staticmethod
    def create_template(user, organization, title="Access template", department=None):
        return DocumentTemplate.objects.create(
            organization=organization,
            department=department,
            created_by=user,
            title=title,
            body_template="Hello {{ full_name }}",
            variables=["full_name"],
        )

    @staticmethod
    def create_document(user, organization, template, title="Access document"):
        return Document.objects.create(
            organization=organization,
            department=template.department,
            template=template,
            created_by=user,
            title=title,
            status=Document.Status.DRAFT,
        )
